"""
services/reporte_microcistina_service.py
Genera el "Reporte de Ensayo" oficial (.docx) de Microcistina LR por campaña,
rellenando la plantilla original (templates/reporte_microcistina_template.docx)
para conservar membrete, logos, firmas y formato del laboratorio.

Toma los datos calculados por el motor ELISA (services/microcistina_service.py)
y los vuelca en la plantilla:
    - Encabezado: código de reporte correlativo.
    - Bloque solicitante: razón social, proyecto, fechas, etc.
    - I. Información: tabla de puntos con coordenadas UTM.
    - II. Resultados: una tabla por cada 3 puntos (concentración en mg/L).
    - III. Control de calidad: control QCS y %CV de duplicados por muestra.

Función pública:
    generar_docx_microcistina_campana(campana_id, overrides) -> bytes
"""

from __future__ import annotations

import copy
import unicodedata
from datetime import date, datetime
from io import BytesIO
from typing import Optional

from docx import Document

from database.client import get_db
from services.microcistina_service import (
    get_codigo_reporte,
    get_corrida,
    get_muestras_microcistina,
    get_param_microcistina,
)

RUTA_PLANTILLA = "templates/reporte_microcistina_template.docx"

_MESES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

DEFAULTS = {
    "razon_social": "AUTORIDAD AUTÓNOMA DE MAJES - AUTODEMA",
    "domicilio": "Urb. La Marina E-8, Cayma - Arequipa - Arequipa",
    "solicitado_por": "Subgerencia de Operación y Mantenimiento",
    "muestreado_por": "Personal LVCA",
    "condicion": "6 °C",
    "tipo_muestra": "Agua para uso y consumo Humano",
    "lugar_emision": "Arequipa",
}


# ─────────────────────────────────────────────────────────────────────────────
# Helpers de formato
# ─────────────────────────────────────────────────────────────────────────────

def _norm(s: str) -> str:
    """Mayúsculas sin acentos para comparar etiquetas."""
    s = unicodedata.normalize("NFKD", s or "")
    s = "".join(c for c in s if not unicodedata.combining(c))
    return s.strip().upper()


def _fmt_num(x: Optional[float], dec: int = 6) -> str:
    """Número con hasta ``dec`` decimales, sin ceros finales sobrantes."""
    if x is None:
        return ""
    s = f"{float(x):.{dec}f}".rstrip("0").rstrip(".")
    return s if s not in ("", "-0") else "0"


def _fmt_fecha(valor) -> str:
    """ISO (YYYY-MM-DD[...]) → dd/mm/YYYY. Acepta date o str."""
    if not valor:
        return ""
    if isinstance(valor, (date, datetime)):
        return valor.strftime("%d/%m/%Y")
    s = str(valor)[:10]
    try:
        return datetime.strptime(s, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return s


def _fecha_larga(valor) -> str:
    """Fecha en formato 'Arequipa, 16 de junio del 2026'."""
    d = None
    if isinstance(valor, (date, datetime)):
        d = valor
    elif valor:
        try:
            d = datetime.strptime(str(valor)[:10], "%Y-%m-%d")
        except ValueError:
            d = None
    if d is None:
        d = datetime.utcnow()
    return f"{d.day} de {_MESES[d.month]} del {d.year}"


def _set_par(paragraph, text: str) -> None:
    """Reemplaza el texto del párrafo conservando el formato del primer run."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _set_cell(cell, text: str) -> None:
    """Escribe en la celda conservando el formato del primer run/párrafo."""
    _set_par(cell.paragraphs[0], text)
    # Elimina párrafos extra de la celda (deja solo el primero).
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


# ─────────────────────────────────────────────────────────────────────────────
# Identificación de tablas de la plantilla
# ─────────────────────────────────────────────────────────────────────────────

def _es_solicitante(t) -> bool:
    return len(t.rows) >= 11 and _norm(t.rows[0].cells[0].text) == "RAZON SOCIAL"


def _es_puntos(t) -> bool:
    return _norm(t.rows[0].cells[0].text).startswith("PUNTOS DE MUESTREO")


def _es_resultados(t) -> bool:
    return (
        len(t.rows) >= 7
        and _norm(t.rows[5].cells[0].text) == "PARAMETROS"
    )


def _es_qc(t) -> bool:
    return (
        len(t.rows) >= 3
        and _norm(t.rows[0].cells[0].text) == "PARAMETRO"
        and _norm(t.rows[0].cells[1].text) == "CONTROL"
    )


# ─────────────────────────────────────────────────────────────────────────────
# Ensamblado de datos
# ─────────────────────────────────────────────────────────────────────────────

def _assemble(campana_id: str, overrides: dict) -> dict:
    db = get_db()
    corrida = get_corrida(campana_id)
    if not corrida or corrida.get("param_a") is None:
        raise ValueError(
            "La campaña no tiene una corrida ELISA de microcistina guardada. "
            "Registra y calcula el ensayo antes de generar el reporte."
        )

    campana = (
        db.table("campanas").select("id, codigo, nombre")
        .eq("id", campana_id).single().execute().data
    ) or {}

    muestras = get_muestras_microcistina(campana_id)
    # Solo muestras con lectura (OD cargada).
    muestras = [m for m in muestras if m.get("od_1") is not None]

    param = get_param_microcistina() or {}
    lmd = param.get("lmd")          # µg/L
    lcm = param.get("lcm")          # µg/L
    lmd_mg = (lmd / 1000.0) if lmd is not None else None
    lcm_mg = (lcm / 1000.0) if lcm is not None else None

    # Fechas auxiliares de las muestras (recepción).
    fecha_recepcion = overrides.get("fecha_recepcion")
    if not fecha_recepcion:
        try:
            ids = [m["id"] for m in muestras]
            rec = (
                db.table("muestras").select("fecha_recepcion_lab")
                .in_("id", ids).execute()
            )
            fechas = sorted(r["fecha_recepcion_lab"] for r in (rec.data or [])
                            if r.get("fecha_recepcion_lab"))
            fecha_recepcion = fechas[0] if fechas else None
        except Exception:
            fecha_recepcion = None

    # Fecha de emisión = validación de resultados; default override o hoy.
    fecha_emision = overrides.get("fecha_emision")
    if not fecha_emision:
        try:
            ids = [m["id"] for m in muestras]
            val = (
                db.table("resultados_laboratorio").select("validado_at")
                .in_("muestra_id", ids).eq("parametro_id", param.get("id"))
                .not_.is_("validado_at", "null").execute()
            )
            vals = sorted(r["validado_at"] for r in (val.data or []) if r.get("validado_at"))
            fecha_emision = vals[-1] if vals else None
        except Exception:
            fecha_emision = None

    # Puntos (Sección I) — únicos por punto.
    puntos = []
    vistos = set()
    for m in muestras:
        p = m.get("punto") or {}
        pid = p.get("id")
        if pid and pid not in vistos:
            vistos.add(pid)
            puntos.append(p)

    # Resultados por muestra (Sección II + %CV Sección III).
    cv_max = float(corrida.get("cv_max") or 15.0)
    filas = []
    for m in muestras:
        p = m.get("punto") or {}
        val_ug = m.get("valor_numerico")
        cualif = m.get("cualificador")
        if cualif == "<LCM" and lcm_mg is not None:
            resultado = f"< {_fmt_num(lcm_mg)}"
        elif val_ug is None:
            resultado = "Fuera de rango"
        else:
            resultado = _fmt_num(val_ug / 1000.0)
        hora = m.get("hora_recoleccion") or ""
        filas.append({
            "punto_codigo": p.get("codigo") or m.get("codigo") or "",
            "descripcion": p.get("descripcion") or p.get("nombre") or "",
            "fecha_hora": f"{_fmt_fecha(m.get('fecha_muestreo'))} {hora}".strip(),
            "tipo_muestra": overrides.get("tipo_muestra") or DEFAULTS["tipo_muestra"],
            "resultado": resultado,
            "cv": m.get("cv_pct"),
        })

    # Procedencia / proyecto por defecto.
    procedencia = overrides.get("procedencia")
    if not procedencia and puntos:
        descs = [p.get("descripcion") or p.get("nombre") for p in puntos if p.get("descripcion") or p.get("nombre")]
        procedencia = max(set(descs), key=descs.count) if descs else campana.get("nombre", "")
    proyecto = overrides.get("proyecto") or campana.get("nombre", "")

    return {
        "corrida": corrida,
        "campana": campana,
        "param": param,
        "lmd_mg": lmd_mg, "lcm_mg": lcm_mg,
        "puntos": puntos,
        "filas": filas,
        "cv_max": cv_max,
        "meta": {
            "codigo_reporte": get_codigo_reporte(campana_id, fecha_emision) or "",
            "razon_social": overrides.get("razon_social") or DEFAULTS["razon_social"],
            "domicilio": overrides.get("domicilio") or DEFAULTS["domicilio"],
            "solicitado_por": overrides.get("solicitado_por") or DEFAULTS["solicitado_por"],
            "proyecto": proyecto,
            "procedencia": procedencia or "",
            "muestreado_por": overrides.get("muestreado_por") or DEFAULTS["muestreado_por"],
            "cantidad": str(len(muestras)),
            "condicion": overrides.get("condicion") or DEFAULTS["condicion"],
            "fecha_recepcion": _fmt_fecha(fecha_recepcion),
            "fecha_ensayo": _fmt_fecha(corrida.get("fecha_ensayo")),
            "fecha_emision": _fmt_fecha(fecha_emision) or _fmt_fecha(datetime.utcnow()),
            "fecha_emision_larga": _fecha_larga(fecha_emision),
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# Rellenado de la plantilla
# ─────────────────────────────────────────────────────────────────────────────

def _fill_solicitante(table, meta: dict) -> None:
    mapa = {
        "RAZON SOCIAL": meta["razon_social"],
        "DOMICILIO LEGAL": meta["domicilio"],
        "SOLICITADO POR": meta["solicitado_por"],
        "PROYECTO": meta["proyecto"],
        "PROCEDENCIA": meta["procedencia"],
        "MUESTREADO POR": meta["muestreado_por"],
        "CANTIDAD DE MUESTRA": meta["cantidad"],
        "CONDICION DE CONSERVACION": meta["condicion"],
        "FECHA DE RECEPCION": meta["fecha_recepcion"],
        "FECHA DE ENSAYO": meta["fecha_ensayo"],
        "FECHA DE EMISION": meta["fecha_emision"],
    }
    for row in table.rows:
        etiqueta = _norm(row.cells[0].text)
        if etiqueta in mapa:
            _set_cell(row.cells[2], mapa[etiqueta])


def _fill_puntos(table, puntos: list[dict]) -> None:
    # Las 2 primeras filas son encabezado; la fila 2 es la plantilla de dato.
    if len(table.rows) < 3:
        return
    plantilla_tr = copy.deepcopy(table.rows[2]._tr)
    # Eliminar filas de datos existentes (de la 2 en adelante).
    for r in list(table.rows)[2:]:
        r._tr.getparent().remove(r._tr)
    for p in puntos:
        nuevo = copy.deepcopy(plantilla_tr)
        table._tbl.append(nuevo)
        celdas = table.rows[-1].cells
        _set_cell(celdas[0], p.get("codigo") or "")
        _set_cell(celdas[1], _fmt_num(p.get("utm_este"), dec=2) or str(p.get("utm_este") or ""))
        _set_cell(celdas[2], _fmt_num(p.get("utm_norte"), dec=2) or str(p.get("utm_norte") or ""))
        _set_cell(celdas[3], p.get("descripcion") or p.get("nombre") or "")
        _set_cell(celdas[4], "-")


def _fill_resultados_tabla(table, grupo: list[dict], lmd_mg, lcm_mg) -> None:
    """Rellena una tabla de resultados con hasta 3 puntos (columnas 4,5,6)."""
    # Filas: 1=estación, 2=descripción, 3=fecha/hora, 4=tipo, 6=resultado.
    for i in range(3):
        col = 4 + i
        if i < len(grupo):
            g = grupo[i]
            _set_cell(table.rows[1].cells[col], g["punto_codigo"])
            _set_cell(table.rows[2].cells[col], g["descripcion"])
            _set_cell(table.rows[3].cells[col], g["fecha_hora"])
            _set_cell(table.rows[4].cells[col], g["tipo_muestra"])
            _set_cell(table.rows[6].cells[col], g["resultado"])
        else:
            for fila in (1, 2, 3, 4, 6):
                _set_cell(table.rows[fila].cells[col], "")
    # Unidad / LDM / LCM en la fila del parámetro (col 1,2,3).
    _set_cell(table.rows[6].cells[1], "mg/L")
    if lmd_mg is not None:
        _set_cell(table.rows[6].cells[2], _fmt_num(lmd_mg))
    if lcm_mg is not None:
        _set_cell(table.rows[6].cells[3], _fmt_num(lcm_mg))


def _fill_qc(table, corrida: dict, cv_max: float) -> None:
    nominal = corrida.get("qcs_nominal")
    tol = corrida.get("qcs_tolerancia")
    c1 = corrida.get("control_conc_1")
    c2 = corrida.get("control_conc_2")
    prom = corrida.get("control_conc_prom")
    cv = corrida.get("control_cv")
    # Rango aceptación y cumplimiento de la concentración del control.
    rango_txt = ""
    cumple_conc = "Si"
    if nominal is not None and tol is not None:
        low, high = float(nominal) - float(tol), float(nominal) + float(tol)
        rango_txt = f"{float(nominal):.3f} ± {float(tol):.3f} ** ({low:.3f} - {high:.3f})"
        if prom is not None:
            cumple_conc = "Si" if low <= float(prom) <= high else "No"
    # Fila 1: concentración del control.
    _set_cell(table.rows[1].cells[1], _fmt_num(c1, 3))
    _set_cell(table.rows[1].cells[2], _fmt_num(c2, 3))
    _set_cell(table.rows[1].cells[3], _fmt_num(prom, 3))
    if rango_txt:
        _set_cell(table.rows[1].cells[4], rango_txt)
    _set_cell(table.rows[1].cells[5], cumple_conc)
    # Fila 2: %CV del control.
    if cv is not None:
        _set_cell(table.rows[2].cells[3], f"{_fmt_num(cv, 2)}%")
    _set_cell(table.rows[2].cells[4], f"≤ {_fmt_num(cv_max, 0)}%")
    cumple_cv = "Si" if (cv is not None and float(cv) <= cv_max) else "No"
    _set_cell(table.rows[2].cells[5], cumple_cv)


def _fill_parrafos(doc, datos: dict) -> None:
    meta = datos["meta"]
    filas = datos["filas"]
    cv_max = datos["cv_max"]
    lote = datos["corrida"].get("kit_lote") or "ABRAXIS/SAES"

    # %CV de duplicados por muestra.
    partes = [f"{f['punto_codigo']}: {_fmt_num(f['cv'], 2)}%" for f in filas if f.get("cv") is not None]
    cv_linea = (
        "Coeficiente de variación (%CV) de duplicados por muestra "
        f"(criterio del fabricante ≤ {_fmt_num(cv_max, 0)}%): " + "; ".join(partes) + "."
    )
    # Observación de muestras fuera de criterio.
    fuera = [f for f in filas if f.get("cv") is not None and float(f["cv"]) > cv_max]
    if fuera:
        lst = "; ".join(f"{f['punto_codigo']} ({_fmt_num(f['cv'], 2)}%)" for f in fuera)
        obs_linea = (
            "Observación: la(s) siguiente(s) muestra(s) presentó(aron) %CV de duplicados "
            f"superior al criterio de ≤ {_fmt_num(cv_max, 0)}% establecido por el fabricante "
            f"({lote}): {lst}. El resultado se reporta a título informativo; se recomienda "
            "su reanálisis para confirmación."
        )
    else:
        obs_linea = (
            "Observación: todas las muestras presentaron %CV de duplicados dentro del "
            f"criterio establecido por el fabricante (≤ {_fmt_num(cv_max, 0)}%)."
        )

    for p in doc.paragraphs:
        t = _norm(p.text)
        if t.startswith("PROYECTO") and ":" in p.text:
            _set_par(p, f"PROYECTO        :        {meta['proyecto']}")
        elif t.startswith("COEFICIENTE DE VARIACION"):
            _set_par(p, cv_linea)
        elif t.startswith("OBSERVACION"):
            _set_par(p, obs_linea)
        elif t.startswith("AREQUIPA,") or t.startswith("AREQUIPA "):
            _set_par(p, f"{DEFAULTS['lugar_emision']}, {meta['fecha_emision_larga']}")


def _set_codigo_reporte(doc, codigo: str) -> None:
    """Reemplaza el código en cada bloque de membrete (runs con 'MC-')."""
    if not codigo:
        return
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if "MC-" in run.text and ("LVCA" in run.text or "-" in run.text):
                            run.text = codigo


# ─────────────────────────────────────────────────────────────────────────────
# API pública
# ─────────────────────────────────────────────────────────────────────────────

def generar_docx_microcistina_campana(
    campana_id: str,
    overrides: Optional[dict] = None,
) -> bytes:
    """Genera el Reporte de Ensayo de microcistina (.docx) y lo devuelve en bytes."""
    overrides = overrides or {}
    datos = _assemble(campana_id, overrides)

    doc = Document(RUTA_PLANTILLA)

    # Código de reporte en los membretes.
    _set_codigo_reporte(doc, datos["meta"]["codigo_reporte"])

    # Localizar tablas.
    resultados_tablas = []
    for t in doc.tables:
        if _es_solicitante(t):
            _fill_solicitante(t, datos["meta"])
        elif _es_puntos(t):
            _fill_puntos(t, datos["puntos"])
        elif _es_resultados(t):
            resultados_tablas.append(t)
        elif _es_qc(t):
            _fill_qc(t, datos["corrida"], datos["cv_max"])

    # Resultados: agrupar puntos de 3 en 3, clonando tablas según haga falta.
    grupos = [datos["filas"][i:i + 3] for i in range(0, len(datos["filas"]), 3)] or [[]]
    if resultados_tablas:
        base = resultados_tablas[0]
        # Asegurar suficientes tablas: clonar la base para grupos extra.
        while len(resultados_tablas) < len(grupos):
            clon = copy.deepcopy(base._tbl)
            resultados_tablas[-1]._tbl.addnext(clon)
            from docx.table import Table
            resultados_tablas.append(Table(clon, base._parent))
        # Eliminar tablas sobrantes.
        for sobra in resultados_tablas[len(grupos):]:
            sobra._tbl.getparent().remove(sobra._tbl)
        resultados_tablas = resultados_tablas[:len(grupos)]
        for t, grupo in zip(resultados_tablas, grupos):
            _fill_resultados_tabla(t, grupo, datos["lmd_mg"], datos["lcm_mg"])

    # Párrafos de cuerpo (proyecto, %CV, observación, fecha de emisión).
    _fill_parrafos(doc, datos)

    salida = BytesIO()
    doc.save(salida)
    return salida.getvalue()
