"""
services/reporte_hidrobiologico_service.py
Genera el reporte hidrobiológico (.docx) por campaña, replicando el formato
de la tabla del laboratorio:
    - Filas de detalle: una fila por especie con el conteo bruto por punto.
    - Filas de resumen por phylum: TOTAL (suma de conteos brutos),
      N° Cel/mL (suma de cel/mL equivalente) y N° Cel/L (= cel/mL × 1000).

Reglas de filtrado solicitadas por el laboratorio:
    - Phyla con <= 3 especies en la taxonomía: muestran TODAS sus especies
      aunque tengan conteo cero en todos los puntos (excepción).
    - Phyla con > 3 especies: muestran sólo las especies con conteo > 0 en
      al menos un punto de la campaña.

Sólo se incluyen puntos cuyas muestras tengan análisis fitoplancton guardado
(``muestras.datos_fitoplancton`` no nulo). Si un punto tiene varias muestras
con análisis en la misma campaña se usa la más reciente.
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from database.client import get_admin_client
from services.fitoplancton_service import TAXONOMIA_FITOPLANCTON


# Phyla con <= UMBRAL_PHYLA_SIN_FILTRAR especies en la taxonomía muestran todas
# sus especies aunque tengan conteo cero. Hoy la regla es "1, 2 o 3 especies".
UMBRAL_PHYLA_SIN_FILTRAR: int = 3


# ─────────────────────────────────────────────────────────────────────────────
# Helpers de formato docx (sombreado, bordes, texto)
# ─────────────────────────────────────────────────────────────────────────────

def _shade_cell(cell, color_hex: str) -> None:
    """Aplica un fondo de color a la celda (color hex sin '#')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _apply_table_borders(table) -> None:
    """Bordes simples grises en todas las aristas de la tabla."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "808080")
        borders.append(b)
    tbl_pr.append(borders)


def _write_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: int = WD_PARAGRAPH_ALIGNMENT.LEFT,
    font_size: float = 9,
    font_color: Optional[str] = None,
) -> None:
    """Escribe texto en una celda con formato controlado."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if font_color:
        r = int(font_color[0:2], 16)
        g = int(font_color[2:4], 16)
        b = int(font_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ─────────────────────────────────────────────────────────────────────────────
# Carga de datos de la campaña
# ─────────────────────────────────────────────────────────────────────────────

def _cargar_muestras_campana(campana_id: str) -> tuple[dict, list[dict]]:
    """
    Devuelve (campana, muestras_unicas_por_punto). Ordena por código de punto.
    Sólo incluye muestras con ``datos_fitoplancton`` no nulo.
    """
    db = get_admin_client()

    campana_res = (
        db.table("campanas")
        .select("id, codigo, nombre, fecha_inicio, fecha_fin")
        .eq("id", campana_id)
        .single()
        .execute()
    )
    campana = campana_res.data or {}

    res = (
        db.table("muestras")
        .select(
            "id, codigo, fecha_muestreo, punto_muestreo_id, datos_fitoplancton, "
            "puntos_muestreo(codigo, nombre)"
        )
        .eq("campana_id", campana_id)
        .not_.is_("datos_fitoplancton", "null")
        .order("fecha_muestreo", desc=False)
        .execute()
    )
    muestras = res.data or []

    # Una muestra por punto (la más reciente si hay duplicados).
    por_punto: dict[str, dict] = {}
    for m in muestras:
        pid = m.get("punto_muestreo_id")
        if not pid:
            continue
        anterior = por_punto.get(pid)
        if anterior is None or (m.get("fecha_muestreo") or "") >= (
            anterior.get("fecha_muestreo") or ""
        ):
            por_punto[pid] = m

    muestras_unicas = sorted(
        por_punto.values(),
        key=lambda m: ((m.get("puntos_muestreo") or {}).get("codigo") or ""),
    )
    return campana, muestras_unicas


def tiene_analisis_hidrobiologico(campana_id: str) -> bool:
    """True si la campaña tiene al menos una muestra con análisis fitoplancton."""
    db = get_admin_client()
    res = (
        db.table("muestras")
        .select("id", count="exact")
        .eq("campana_id", campana_id)
        .not_.is_("datos_fitoplancton", "null")
        .limit(1)
        .execute()
    )
    return bool(res.data) or bool(getattr(res, "count", 0) or 0)


# ─────────────────────────────────────────────────────────────────────────────
# Agregaciones por phylum
# ─────────────────────────────────────────────────────────────────────────────

def _construir_matrices(
    muestras: list[dict],
) -> tuple[dict[str, dict[str, list[int]]], dict[str, list[dict[str, float]]]]:
    """
    Devuelve dos estructuras:
        conteos[phylum][especie] = [conteo_bruto_punto_0, ..., conteo_bruto_punto_N]
        agregados[phylum][i]    = {"total": int, "cel_ml": float}

    Las especies que aparezcan en el JSONB pero no estén en
    ``TAXONOMIA_FITOPLANCTON`` (datos legacy / renombrados) se ignoran en
    las filas de detalle, pero SÍ se suman al total y al cel/mL del phylum
    para que el resumen refleje el dato real cargado.
    """
    n = len(muestras)
    conteos: dict[str, dict[str, list[int]]] = {}
    agregados: dict[str, list[dict[str, float]]] = {}

    for filo, especies_def in TAXONOMIA_FITOPLANCTON.items():
        conteos[filo] = {e["nombre"]: [0] * n for e in especies_def}
        agregados[filo] = [{"total": 0, "cel_ml": 0.0} for _ in range(n)]

    for idx, muestra in enumerate(muestras):
        doc = muestra.get("datos_fitoplancton") or {}
        resultados = doc.get("resultados") or {}
        for filo, especies_data in resultados.items():
            if filo not in agregados:
                # Phylum legacy ya no presente en la taxonomía. Lo ignoramos
                # (caso raro porque el JSONB se reescribió en abril 2026).
                continue
            for esp_nombre, val in (especies_data or {}).items():
                conteo_bruto = int(val.get("conteo_bruto") or 0)
                cel_ml_equiv = float(
                    val.get("cel_ml_equiv")
                    if val.get("cel_ml_equiv") is not None
                    else val.get("cel_ml") or 0.0
                )
                if esp_nombre in conteos[filo]:
                    conteos[filo][esp_nombre][idx] += conteo_bruto
                agregados[filo][idx]["total"] += conteo_bruto
                agregados[filo][idx]["cel_ml"] += cel_ml_equiv

    return conteos, agregados


def _especies_visibles_por_phylum(
    conteos: dict[str, dict[str, list[int]]],
) -> dict[str, list[str]]:
    """
    Aplica la regla:
      - Phyla con <= UMBRAL especies: todas visibles aunque sean cero.
      - Phyla con > UMBRAL especies: sólo las con conteo > 0 en >=1 punto.
    """
    visibles: dict[str, list[str]] = {}
    for filo, especies_def in TAXONOMIA_FITOPLANCTON.items():
        nombres = [e["nombre"] for e in especies_def]
        if len(nombres) <= UMBRAL_PHYLA_SIN_FILTRAR:
            visibles[filo] = nombres
        else:
            visibles[filo] = [
                n for n in nombres
                if any((c or 0) > 0 for c in conteos[filo][n])
            ]
    return visibles


# ─────────────────────────────────────────────────────────────────────────────
# Generador del .docx
# ─────────────────────────────────────────────────────────────────────────────

def generar_docx_hidrobiologico_campana(campana_id: str) -> bytes:
    """
    Genera el documento Word (.docx) con la tabla hidrobiológica de la campaña.

    Lanza ``ValueError`` si la campaña no tiene ninguna muestra con análisis
    fitoplancton guardado.
    """
    campana, muestras = _cargar_muestras_campana(campana_id)
    if not muestras:
        raise ValueError(
            "La campaña no tiene análisis hidrobiológicos guardados. "
            "Carga al menos una muestra con datos de fitoplancton antes de descargar."
        )

    puntos_codigos: list[str] = [
        ((m.get("puntos_muestreo") or {}).get("codigo") or "—") for m in muestras
    ]
    n_puntos = len(muestras)

    conteos, agregados = _construir_matrices(muestras)
    visibles = _especies_visibles_por_phylum(conteos)

    # ── Construir documento ────────────────────────────────────────────────
    document = Document()

    # Página A4 horizontal con márgenes ajustados.
    section = document.sections[0]
    if section.orientation != WD_ORIENTATION.LANDSCAPE:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Título: "Resultados de abundancia de fitoplancton — <nombre campaña>"
    nombre_campana = (campana.get("nombre") or campana.get("codigo") or "").strip()
    titulo = document.add_paragraph()
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titulo.paragraph_format.space_after = Pt(6)
    run_t = titulo.add_run(
        f"Resultados de abundancia de fitoplancton — {nombre_campana}"
    )
    run_t.font.name = "Arial"
    run_t.font.size = Pt(13)
    run_t.bold = True

    # Subtítulo con código y rango de fechas (si existen).
    fi = campana.get("fecha_inicio") or ""
    ff = campana.get("fecha_fin") or ""
    sub_partes = []
    if campana.get("codigo"):
        sub_partes.append(f"Código: {campana['codigo']}")
    if fi or ff:
        sub_partes.append(f"Periodo: {fi}" + (f" → {ff}" if ff and ff != fi else ""))
    sub_partes.append(f"Puntos analizados: {n_puntos}")
    if sub_partes:
        sub = document.add_paragraph()
        sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub.paragraph_format.space_after = Pt(8)
        run_s = sub.add_run(" · ".join(sub_partes))
        run_s.font.name = "Arial"
        run_s.font.size = Pt(9)
        run_s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Tabla principal: 2 cols (Phylum + Especie) + 1 col por punto
    n_cols = 2 + n_puntos
    table = document.add_table(rows=1, cols=n_cols)
    table.autofit = True
    _apply_table_borders(table)

    # Header
    hdr = table.rows[0]
    hdr_phylum_sp = hdr.cells[0].merge(hdr.cells[1])
    _write_cell(
        hdr_phylum_sp,
        "Phylum / sp.",
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        font_color="FFFFFF",
    )
    _shade_cell(hdr_phylum_sp, "404040")
    for i, code in enumerate(puntos_codigos):
        c = hdr.cells[2 + i]
        _write_cell(
            c,
            code,
            bold=True,
            align=WD_PARAGRAPH_ALIGNMENT.CENTER,
            font_color="FFFFFF",
            font_size=8.5,
        )
        _shade_cell(c, "404040")

    # Filas por phylum
    color_phylum_bg = "DCE6F1"   # azul clarito para la columna phylum
    color_resumen_bg = "F2F2F2"  # gris para filas TOTAL/Cel/mL/Cel/L

    for filo, especies_def in TAXONOMIA_FITOPLANCTON.items():
        especies_a_mostrar = visibles.get(filo, [])
        # Phyla con > UMBRAL y sin especies visibles ni totales: omitir bloque entero.
        total_phylum = sum(a["total"] for a in agregados[filo])
        if (
            not especies_a_mostrar
            and len(especies_def) > UMBRAL_PHYLA_SIN_FILTRAR
            and total_phylum == 0
        ):
            continue

        # Filas de detalle
        for esp_nombre in especies_a_mostrar:
            row = table.add_row()
            _write_cell(
                row.cells[0],
                filo.upper(),
                bold=True,
                align=WD_PARAGRAPH_ALIGNMENT.LEFT,
                font_size=8.5,
            )
            _shade_cell(row.cells[0], color_phylum_bg)
            _write_cell(
                row.cells[1],
                esp_nombre,
                italic=True,
                align=WD_PARAGRAPH_ALIGNMENT.LEFT,
                font_size=9,
            )
            for i in range(n_puntos):
                val = conteos[filo][esp_nombre][i]
                _write_cell(
                    row.cells[2 + i],
                    str(int(val)),
                    align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    font_size=9,
                )

        # 3 filas de resumen del phylum
        for label, key, factor in (
            ("TOTAL",     "total",  1),
            ("N° Cel/mL", "cel_ml", 1),
            ("N° Cel/L",  "cel_ml", 1000),
        ):
            row = table.add_row()
            etiqueta = row.cells[0].merge(row.cells[1])
            _write_cell(
                etiqueta,
                label,
                bold=True,
                align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
                font_size=9,
            )
            _shade_cell(etiqueta, color_resumen_bg)
            for i in range(n_puntos):
                ag = agregados[filo][i]
                if key == "total":
                    val_int = int(ag["total"])
                else:
                    val_int = int(round(ag["cel_ml"] * factor))
                c = row.cells[2 + i]
                _write_cell(
                    c,
                    f"{val_int:,}".replace(",", " "),
                    bold=True,
                    align=WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    font_size=9,
                )
                _shade_cell(c, color_resumen_bg)

    # Pie con timestamp.
    pie = document.add_paragraph()
    pie.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pie.paragraph_format.space_before = Pt(8)
    run_p = pie.add_run(
        f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    run_p.font.name = "Arial"
    run_p.font.size = Pt(8)
    run_p.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
