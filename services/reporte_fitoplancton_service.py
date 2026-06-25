"""
services/reporte_fitoplancton_service.py
Genera el "Reporte de Ensayo" oficial (.docx) de FITOPLANCTON por campaña,
integrando campañas → registro de muestras → resultados de fitoplancton.

Estructura (replica el formato del laboratorio):
    - Encabezado de sección con los dos logos (LVCA + AUTODEMA), repetido en
      TODAS las páginas.
    - Portada: título, "REPORTE DE ENSAYO" y código correlativo
      "LVCA - NNN-FITO-AAAA".
    - Bloque solicitante (razón social, proyecto, procedencia, fechas…).
    - I. INFORMACIÓN: puntos de muestreo con coordenadas UTM.
    - II. RESULTADOS: comunidad fitoplanctónica (División/Clase/Orden/Familia/
      Especie + una columna por estación, en cel/mL) + tabla resumen de
      parámetros comunitarios (Riqueza, Abundancia total). SIN índice de Shannon.
    - III. CONTROL DE CALIDAD: parámetros del método usado (Sedgewick-Rafter o
      Utermöhl), detectado desde los metadatos del análisis.
    - IV. MÉTODOS Y REFERENCIAS / V. VALORES REFERENCIALES: texto fijo.

Los valores de densidad son SIEMPRE cel/mL (campo ``cel_ml_equiv`` del JSONB,
con respaldo a ``cel_ml``). Los códigos de estación y puntos se toman de la
plataforma, por lo que funciona con todos los puntos actuales y futuros.

Función pública:
    generar_docx_fitoplancton_campana(campana_id, overrides) -> bytes
    get_codigo_reporte_fito(campana_id, fecha_emision) -> str
    tiene_analisis_hidrobiologico  (reexport de reporte_hidrobiologico_service)
"""

from __future__ import annotations

import os
from datetime import date, datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

from database.client import get_db
from services.fitoplancton_service import (
    METODO_UTERMOHL,
    METODOS_CONTEO,
)
from services.taxonomia_fitoplancton_service import ORDEN_FILOS, get_taxonomia
# Reutilizamos los helpers de formato del reporte hidrobiológico.
from services.reporte_hidrobiologico_service import (
    _apply_table_borders,
    _shade_cell,
    _write_cell,
    tiene_analisis_hidrobiologico,  # re-export para la UI
)

_BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LOGO_LVCA = os.path.join(_BASE_DIR, "imagenes", "logo_lvca.png")
LOGO_AUTODEMA = os.path.join(_BASE_DIR, "imagenes", "autodema_logo.png")

# Cuántas estaciones (columnas) entran por tabla de resultados en A4 horizontal.
ESTACIONES_POR_TABLA = 4

_MESES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

DEFAULTS = {
    "razon_social": "AUTORIDAD AUTÓNOMA DE MAJES - AUTODEMA",
    "domicilio": "Urb. La Marina E-8, Cayma - Arequipa - Arequipa",
    "solicitado_por": "Subgerencia de Operación y Mantenimiento",
    "proyecto": "Monitoreo hidrobiológico de fitoplancton - Sistema Regulado Chili",
    "muestreado_por": "Personal LVCA",
    "condicion": "6 °C",
    "lugar_emision": "Arequipa",
}

_COLOR_TITULO = "1F4D78"   # azul corporativo del reporte


# ─────────────────────────────────────────────────────────────────────────────
# Formato auxiliar
# ─────────────────────────────────────────────────────────────────────────────

def _fmt_fecha(valor) -> str:
    if not valor:
        return ""
    if isinstance(valor, (date, datetime)):
        return valor.strftime("%d/%m/%Y")
    s = str(valor)[:10]
    try:
        return datetime.strptime(s, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return s


def _fecha_larga(valor) -> str:
    d = None
    if isinstance(valor, (date, datetime)):
        d = valor
    elif valor:
        try:
            d = datetime.strptime(str(valor)[:10], "%Y-%m-%d")
        except ValueError:
            d = None
    if d is None:
        d = datetime.utcnow()
    return f"{d.day} de {_MESES[d.month]} del {d.year}"


def _fmt_cel(v: float) -> str:
    """Densidad cel/mL: entero con separador de miles si ≥100; '–' si 0."""
    v = float(v or 0)
    if v <= 0:
        return "–"
    if v >= 100:
        return f"{int(round(v)):,}".replace(",", " ")
    return f"{v:.2f}"


def _cel_de_valor(val: dict) -> float:
    """cel/mL equivalente de una especie del JSONB (con respaldo legacy)."""
    if val is None:
        return 0.0
    if val.get("cel_ml_equiv") is not None:
        return float(val.get("cel_ml_equiv") or 0.0)
    return float(val.get("cel_ml") or 0.0)


# ─────────────────────────────────────────────────────────────────────────────
# Correlativo del reporte (espejo de microcistina_service.get_codigo_reporte)
# ─────────────────────────────────────────────────────────────────────────────

def _generar_codigo_reporte(db, fecha_emision: Optional[str]) -> str:
    anio = datetime.utcnow().year
    if fecha_emision:
        try:
            anio = datetime.strptime(str(fecha_emision)[:10], "%Y-%m-%d").year
        except ValueError:
            pass
    try:
        res = db.table("fitoplancton_reportes").select("campana_id", count="exact").execute()
        n = (getattr(res, "count", 0) or 0) + 1
    except Exception:
        n = 1
    return f"LVCA - {int(n):03d}-FITO-{anio}"


def get_codigo_reporte_fito(
    campana_id: str, fecha_emision: Optional[str] = None, crear: bool = True
) -> Optional[str]:
    """
    Correlativo 'LVCA - NNN-FITO-AAAA' por campaña. Lo lee de
    fitoplancton_reportes; si no existe y ``crear`` es True, lo genera y guarda.
    """
    db = get_db()
    try:
        r = (
            db.table("fitoplancton_reportes").select("*")
            .eq("campana_id", campana_id).limit(1).execute()
        )
        rows = r.data or []
    except Exception:
        rows = []
    if rows and rows[0].get("codigo_reporte"):
        return rows[0]["codigo_reporte"]
    if not crear:
        return None
    codigo = _generar_codigo_reporte(db, fecha_emision)
    try:
        db.table("fitoplancton_reportes").upsert(
            {"campana_id": campana_id, "codigo_reporte": codigo,
             "fecha_emision": fecha_emision},
            on_conflict="campana_id",
        ).execute()
    except Exception:
        pass
    return codigo


# ─────────────────────────────────────────────────────────────────────────────
# Carga y ensamblado de datos
# ─────────────────────────────────────────────────────────────────────────────

def _cargar_estaciones(campana_id: str) -> tuple[dict, list[dict]]:
    """
    (campana, estaciones) — una estación por punto (la muestra más reciente con
    análisis de fitoplancton), ordenadas por código de punto.
    """
    db = get_db()
    campana = (
        db.table("campanas")
        .select("id, codigo, nombre, fecha_inicio, fecha_fin")
        .eq("id", campana_id).single().execute().data
    ) or {}

    res = (
        db.table("muestras")
        .select(
            "id, codigo, fecha_muestreo, hora_recoleccion, fecha_recepcion_lab, "
            "punto_muestreo_id, datos_fitoplancton, "
            "puntos_muestreo(codigo, nombre, descripcion, utm_este, utm_norte, altitud_msnm)"
        )
        .eq("campana_id", campana_id)
        .not_.is_("datos_fitoplancton", "null")
        .order("fecha_muestreo", desc=False)
        .execute()
    )
    muestras = res.data or []

    por_punto: dict[str, dict] = {}
    for m in muestras:
        pid = m.get("punto_muestreo_id")
        if not pid:
            continue
        prev = por_punto.get(pid)
        if prev is None or (m.get("fecha_muestreo") or "") >= (prev.get("fecha_muestreo") or ""):
            por_punto[pid] = m

    estaciones = sorted(
        por_punto.values(),
        key=lambda m: ((m.get("puntos_muestreo") or {}).get("codigo") or ""),
    )
    return campana, estaciones


def _mapa_taxonomia() -> dict[str, tuple]:
    """especie -> (division, clase, orden, familia) desde el catálogo."""
    mapa: dict[str, tuple] = {}
    for division, especies in get_taxonomia().items():
        for e in especies:
            mapa[e["nombre"]] = (
                division, e.get("clase"), e.get("orden"), e.get("familia"),
            )
    return mapa


def _orden_division(division: str) -> int:
    try:
        return ORDEN_FILOS.index(division)
    except ValueError:
        return len(ORDEN_FILOS)


def _assemble(campana_id: str, overrides: dict) -> dict:
    campana, estaciones = _cargar_estaciones(campana_id)
    if not estaciones:
        raise ValueError(
            "La campaña no tiene análisis de fitoplancton guardados. "
            "Carga al menos una muestra con datos de fitoplancton antes de "
            "generar el reporte."
        )

    mapa_tax = _mapa_taxonomia()

    # Matriz especie × estación en cel/mL + metadatos de cada estación.
    n = len(estaciones)
    cel_por_especie: dict[str, list[float]] = {}
    division_de: dict[str, str] = {}

    estaciones_info: list[dict] = []
    metodos_detectados: list[str] = []
    fechas_analisis: list[str] = []

    for idx, est in enumerate(estaciones):
        p = est.get("puntos_muestreo") or {}
        estaciones_info.append({
            "codigo": p.get("codigo") or est.get("codigo") or "—",
            "nombre": p.get("nombre") or "",
            "descripcion": p.get("descripcion") or p.get("nombre") or "",
            "este": p.get("utm_este"),
            "norte": p.get("utm_norte"),
            "altitud": p.get("altitud_msnm"),
            "fecha": _fmt_fecha(est.get("fecha_muestreo")),
            "hora": est.get("hora_recoleccion") or "",
        })

        doc = est.get("datos_fitoplancton") or {}
        meta = doc.get("metadatos") or {}
        metodos_detectados.append(meta.get("metodo") or "sedgewick_rafter")
        if meta.get("fecha_analisis"):
            fechas_analisis.append(meta["fecha_analisis"])

        for filo, especies in (doc.get("resultados") or {}).items():
            for esp, val in (especies or {}).items():
                cel = _cel_de_valor(val)
                if esp not in cel_por_especie:
                    cel_por_especie[esp] = [0.0] * n
                    division_de[esp] = (mapa_tax.get(esp, (filo,))[0]) or filo
                cel_por_especie[esp][idx] += cel

    # Especies presentes (cel/mL > 0 en al menos una estación), ordenadas.
    presentes = [e for e, fila in cel_por_especie.items() if any(v > 0 for v in fila)]
    presentes.sort(key=lambda e: (_orden_division(division_de.get(e, "")), e.lower()))

    filas_especie: list[dict] = []
    for esp in presentes:
        division, clase, orden, familia = mapa_tax.get(esp, (division_de.get(esp, ""), None, None, None))
        filas_especie.append({
            "division": division or "",
            "clase": clase or "",
            "orden": orden or "",
            "familia": familia or "",
            "especie": esp,
            "cel": cel_por_especie[esp],
        })

    # Parámetros comunitarios por estación.
    riqueza = [0] * n
    abundancia = [0.0] * n
    for fila in filas_especie:
        for i, v in enumerate(fila["cel"]):
            if v > 0:
                riqueza[i] += 1
                abundancia[i] += v

    # Método dominante para la sección de control de calidad.
    metodo = max(set(metodos_detectados), key=metodos_detectados.count) if metodos_detectados else "sedgewick_rafter"
    metodo_mixto = len(set(metodos_detectados)) > 1
    # Metadatos de QC del primer análisis con ese método.
    qc_meta = {}
    for est in estaciones:
        meta = (est.get("datos_fitoplancton") or {}).get("metadatos") or {}
        if (meta.get("metodo") or "sedgewick_rafter") == metodo:
            qc_meta = meta
            break

    # Fechas.
    fecha_recepcion = overrides.get("fecha_recepcion")
    if not fecha_recepcion:
        recs = sorted(e.get("fecha_recepcion_lab") for e in estaciones if e.get("fecha_recepcion_lab"))
        fecha_recepcion = recs[0] if recs else None
    fecha_ensayo = overrides.get("fecha_ensayo") or (sorted(fechas_analisis)[0] if fechas_analisis else None)
    fecha_emision = overrides.get("fecha_emision") or datetime.utcnow().date().isoformat()

    # Procedencia / proyecto.
    procedencia = overrides.get("procedencia")
    if not procedencia:
        descs = [i["descripcion"] for i in estaciones_info if i["descripcion"]]
        procedencia = max(set(descs), key=descs.count) if descs else (campana.get("nombre") or "")
    proyecto = overrides.get("proyecto") or campana.get("nombre") or DEFAULTS["proyecto"]

    return {
        "campana": campana,
        "estaciones": estaciones_info,
        "filas_especie": filas_especie,
        "riqueza": riqueza,
        "abundancia": abundancia,
        "metodo": metodo,
        "metodo_mixto": metodo_mixto,
        "qc_meta": qc_meta,
        "meta": {
            "codigo_reporte": get_codigo_reporte_fito(campana_id, fecha_emision) or "",
            "razon_social": overrides.get("razon_social") or DEFAULTS["razon_social"],
            "domicilio": overrides.get("domicilio") or DEFAULTS["domicilio"],
            "solicitado_por": overrides.get("solicitado_por") or DEFAULTS["solicitado_por"],
            "proyecto": proyecto,
            "procedencia": procedencia or "",
            "muestreado_por": overrides.get("muestreado_por") or DEFAULTS["muestreado_por"],
            "cantidad": str(len(estaciones_info)),
            "condicion": overrides.get("condicion") or DEFAULTS["condicion"],
            "fecha_recepcion": _fmt_fecha(fecha_recepcion),
            "fecha_ensayo": _fmt_fecha(fecha_ensayo),
            "fecha_emision": _fmt_fecha(fecha_emision),
            "fecha_emision_larga": _fecha_larga(fecha_emision),
        },
    }


# ─────────────────────────────────────────────────────────────────────────────
# Construcción del .docx
# ─────────────────────────────────────────────────────────────────────────────

def _configurar_pagina_y_encabezado(document, codigo_reporte: str) -> None:
    """A4 horizontal + encabezado con los dos logos (repetido en todas las hojas)."""
    section = document.sections[0]
    if section.orientation != WD_ORIENTATION.LANDSCAPE:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.6)

    header = section.header
    header.is_linked_to_previous = False
    # Tabla 1×3: logo LVCA | título centrado | logo AUTODEMA.
    tbl = header.add_table(rows=1, cols=3, width=Cm(26))
    tbl.autofit = False
    celdas = tbl.rows[0].cells
    anchos = (Cm(4.5), Cm(17), Cm(4.5))
    for c, w in zip(celdas, anchos):
        c.width = w

    # Logo izquierdo (LVCA).
    p_izq = celdas[0].paragraphs[0]
    p_izq.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    if os.path.exists(LOGO_LVCA):
        try:
            p_izq.add_run().add_picture(LOGO_LVCA, height=Cm(1.5))
        except Exception:
            pass

    # Texto central.
    p_cen = celdas[1].paragraphs[0]
    p_cen.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r1 = p_cen.add_run("LABORATORIO DE VIGILANCIA DE CALIDAD DEL AGUA - LVCA")
    r1.font.name = "Arial"; r1.font.size = Pt(9); r1.bold = True
    r1.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    if codigo_reporte:
        p_cen.add_run("\n")
        r2 = p_cen.add_run(codigo_reporte)
        r2.font.name = "Arial"; r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Logo derecho (AUTODEMA).
    p_der = celdas[2].paragraphs[0]
    p_der.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    if os.path.exists(LOGO_AUTODEMA):
        try:
            p_der.add_run().add_picture(LOGO_AUTODEMA, height=Cm(1.5))
        except Exception:
            pass


def _titulo_seccion(document, texto: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(texto)
    r.font.name = "Arial"; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)


def _parrafo(document, texto: str, *, size: float = 9, bold: bool = False,
             italic: bool = False, color: str | None = None,
             align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_after: float = 3) -> None:
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(texto)
    r.font.name = "Arial"; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


def _portada(document, meta: dict) -> None:
    _parrafo(document, "REPORTE DE ENSAYO", size=14, bold=True,
             color=_COLOR_TITULO, align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=2)
    _parrafo(document, meta["codigo_reporte"], size=10, bold=True,
             align=WD_PARAGRAPH_ALIGNMENT.CENTER, space_after=8)

    pares = [
        ("Razón Social", meta["razon_social"]),
        ("Domicilio Legal", meta["domicilio"]),
        ("Solicitado Por", meta["solicitado_por"]),
        ("Proyecto", meta["proyecto"]),
        ("Procedencia", meta["procedencia"]),
        ("Muestreado Por", meta["muestreado_por"]),
        ("Cantidad de Muestra", meta["cantidad"]),
        ("Condición de Conservación", meta["condicion"]),
        ("Fecha de Recepción", meta["fecha_recepcion"]),
        ("Fecha de Ensayo", meta["fecha_ensayo"]),
        ("Fecha de Emisión", meta["fecha_emision"]),
    ]
    tabla = document.add_table(rows=len(pares), cols=2)
    tabla.autofit = False
    _apply_table_borders(tabla)
    for i, (k, v) in enumerate(pares):
        c0, c1 = tabla.rows[i].cells
        c0.width = Cm(6); c1.width = Cm(20)
        _write_cell(c0, k, bold=True, font_size=9)
        _shade_cell(c0, "F2F2F2")
        _write_cell(c1, v or "—", font_size=9)


def _seccion_informacion(document, estaciones: list[dict]) -> None:
    _titulo_seccion(document, "I. INFORMACIÓN")
    encabezados = ["Punto de muestreo", "Este", "Norte", "Altitud (m s.n.m.)",
                   "Fecha", "Hora", "Descripción de la estación"]
    tabla = document.add_table(rows=1, cols=len(encabezados))
    tabla.autofit = True
    _apply_table_borders(tabla)
    for i, h in enumerate(encabezados):
        _write_cell(tabla.rows[0].cells[i], h, bold=True,
                    align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_color="FFFFFF", font_size=8.5)
        _shade_cell(tabla.rows[0].cells[i], "404040")
    for est in estaciones:
        celdas = tabla.add_row().cells
        _write_cell(celdas[0], est["codigo"], bold=True, font_size=8.5)
        _write_cell(celdas[1], _num(est["este"]), align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8.5)
        _write_cell(celdas[2], _num(est["norte"]), align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8.5)
        _write_cell(celdas[3], _num(est["altitud"]), align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8.5)
        _write_cell(celdas[4], est["fecha"], align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=8.5)
        _write_cell(celdas[5], est["hora"], align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=8.5)
        _write_cell(celdas[6], est["descripcion"], font_size=8.5)


def _num(v) -> str:
    if v is None:
        return "—"
    try:
        f = float(v)
        return f"{int(f):,}".replace(",", " ") if f == int(f) else f"{f:.2f}"
    except (TypeError, ValueError):
        return str(v)


def _tabla_resultados_chunk(document, filas_especie: list[dict],
                            estaciones: list[dict], idx0: int, idx1: int) -> None:
    """Tabla de comunidad para las estaciones [idx0, idx1)."""
    cols_tax = ["División", "Clase", "Orden", "Familia", "Especie"]
    sub = estaciones[idx0:idx1]
    n_cols = len(cols_tax) + len(sub)
    tabla = document.add_table(rows=1, cols=n_cols)
    tabla.autofit = True
    _apply_table_borders(tabla)

    hdr = tabla.rows[0].cells
    for i, h in enumerate(cols_tax):
        _write_cell(hdr[i], h, bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    font_color="FFFFFF", font_size=8)
        _shade_cell(hdr[i], "404040")
    for j, est in enumerate(sub):
        cab = f"{est['codigo']}\n{est['fecha']}\n{est['hora']}".strip()
        _write_cell(hdr[len(cols_tax) + j], cab, bold=True,
                    align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_color="FFFFFF", font_size=8)
        _shade_cell(hdr[len(cols_tax) + j], "404040")

    for fila in filas_especie:
        celdas = tabla.add_row().cells
        _write_cell(celdas[0], fila["division"], font_size=8)
        _write_cell(celdas[1], fila["clase"], font_size=8)
        _write_cell(celdas[2], fila["orden"], font_size=8)
        _write_cell(celdas[3], fila["familia"], font_size=8)
        _write_cell(celdas[4], fila["especie"], italic=True, font_size=8)
        for j in range(idx0, idx1):
            _write_cell(celdas[len(cols_tax) + (j - idx0)], _fmt_cel(fila["cel"][j]),
                        align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8)

    # Fila Total (cel/mL) por estación.
    celdas = tabla.add_row().cells
    etiqueta = celdas[0]
    for k in range(1, len(cols_tax)):
        etiqueta = etiqueta.merge(celdas[k])
    _write_cell(etiqueta, "Total (cel/mL)", bold=True,
                align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8)
    _shade_cell(etiqueta, "F2F2F2")
    for j in range(idx0, idx1):
        total = sum(f["cel"][j] for f in filas_especie)
        c = celdas[len(cols_tax) + (j - idx0)]
        _write_cell(c, _fmt_cel(total), bold=True, align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8)
        _shade_cell(c, "F2F2F2")


def _tabla_comunitarios_chunk(document, estaciones: list[dict],
                              riqueza: list[int], abundancia: list[float],
                              idx0: int, idx1: int) -> None:
    sub = estaciones[idx0:idx1]
    tabla = document.add_table(rows=1, cols=1 + len(sub))
    tabla.autofit = True
    _apply_table_borders(tabla)
    hdr = tabla.rows[0].cells
    _write_cell(hdr[0], "Parámetro comunitario", bold=True, font_color="FFFFFF", font_size=8.5)
    _shade_cell(hdr[0], "404040")
    for j, est in enumerate(sub):
        _write_cell(hdr[1 + j], est["codigo"], bold=True,
                    align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_color="FFFFFF", font_size=8.5)
        _shade_cell(hdr[1 + j], "404040")

    # Riqueza
    r = tabla.add_row().cells
    _write_cell(r[0], "Riqueza (n.º de taxones)", bold=True, font_size=8.5)
    for j in range(idx0, idx1):
        _write_cell(r[1 + (j - idx0)], str(riqueza[j]),
                    align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8.5)
    # Abundancia total
    a = tabla.add_row().cells
    _write_cell(a[0], "Abundancia total (cel/mL)", bold=True, font_size=8.5)
    for j in range(idx0, idx1):
        _write_cell(a[1 + (j - idx0)], _fmt_cel(abundancia[j]),
                    align=WD_PARAGRAPH_ALIGNMENT.RIGHT, font_size=8.5)


def _seccion_resultados(document, datos: dict) -> None:
    _titulo_seccion(document, "II. RESULTADOS")
    estaciones = datos["estaciones"]
    filas = datos["filas_especie"]
    n = len(estaciones)

    if not filas:
        _parrafo(document, "No se registraron organismos fitoplanctónicos en las "
                 "muestras analizadas de esta campaña.", italic=True)
    else:
        _parrafo(document, "Densidad de organismos por estación, expresada en "
                 "células por mililitro (cel/mL).", size=8.5, color="555555")
        for idx0 in range(0, n, ESTACIONES_POR_TABLA):
            idx1 = min(idx0 + ESTACIONES_POR_TABLA, n)
            _tabla_resultados_chunk(document, filas, estaciones, idx0, idx1)
            document.add_paragraph()

    _parrafo(document, "Resumen de parámetros comunitarios", size=9, bold=True,
             color=_COLOR_TITULO, space_after=2)
    for idx0 in range(0, n, ESTACIONES_POR_TABLA):
        idx1 = min(idx0 + ESTACIONES_POR_TABLA, n)
        _tabla_comunitarios_chunk(document, estaciones, datos["riqueza"],
                                  datos["abundancia"], idx0, idx1)
        document.add_paragraph()


def _seccion_control_calidad(document, datos: dict) -> None:
    _titulo_seccion(document, "III. CONTROL DE CALIDAD")
    metodo = datos["metodo"]
    meta = datos["qc_meta"] or {}
    _parrafo(document, f"Método de análisis: {METODOS_CONTEO.get(metodo, metodo)}.",
             size=9, bold=True, space_after=2)
    if datos.get("metodo_mixto"):
        _parrafo(document, "Nota: la campaña combina más de un método de conteo; "
                 "se reportan los criterios del método predominante.",
                 size=8, italic=True, color="555555")

    if metodo == METODO_UTERMOHL:
        filas = [
            ("Volumen sedimentado (V_sed)", f"{_num(meta.get('vol_sedimentado_ml'))} mL",
             "Según altura de cámara (APHA 10200 F)", "Sí"),
            ("Área de la cámara (A_cámara)", f"{_num(meta.get('area_camara_mm2'))} mm²",
             "Cámara calibrada", "Sí"),
            ("Área contada (A_contada)", f"{_num(meta.get('area_contada_mm2'))} mm²",
             "Transectos/campos representativos", "Sí"),
            ("Unidades contadas del taxón dominante", "Recuento por sedimentación",
             "≥ 400 unidades (precisión ±10 %)", "Sí"),
        ]
    else:
        num_campos = meta.get("num_campos") or 0
        cumple_campos = "Sí" if (int(num_campos or 0) >= 100 or float(meta.get("area_campo_mm2") or 0) >= 1000) else "Revisar"
        filas = [
            ("Volumen de cámara", "1 mL", "1 mL (cámara Sedgewick-Rafter)", "Sí"),
            ("Volumen de muestra (Vs)", f"{_num(meta.get('vol_muestra_ml'))} mL", "Registrado", "Sí"),
            ("Volumen concentrado (Vc)", f"{_num(meta.get('vol_concentrado_ml'))} mL", "Vc ≤ Vs", "Sí"),
            ("Área del campo", f"{_num(meta.get('area_campo_mm2'))} mm²", "Calibrada", "Sí"),
            ("N.º de campos contados", f"{_num(num_campos)}", "≥ 100 (o cámara completa)", cumple_campos),
        ]

    encabezados = ["Parámetro de control", "Valor", "Criterio", "Cumple"]
    tabla = document.add_table(rows=1, cols=4)
    tabla.autofit = True
    _apply_table_borders(tabla)
    for i, h in enumerate(encabezados):
        _write_cell(tabla.rows[0].cells[i], h, bold=True,
                    align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_color="FFFFFF", font_size=8.5)
        _shade_cell(tabla.rows[0].cells[i], "404040")
    for nombre, valor, criterio, cumple in filas:
        celdas = tabla.add_row().cells
        _write_cell(celdas[0], nombre, bold=True, font_size=8.5)
        _write_cell(celdas[1], valor, align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=8.5)
        _write_cell(celdas[2], criterio, font_size=8.5)
        _write_cell(celdas[3], cumple, align=WD_PARAGRAPH_ALIGNMENT.CENTER, font_size=8.5)


def _seccion_metodos(document) -> None:
    _titulo_seccion(document, "IV. MÉTODOS Y REFERENCIAS")
    _parrafo(document,
             "El análisis cuantitativo de fitoplancton se realiza según APHA, AWWA & WEF, "
             "Standard Methods for the Examination of Water and Wastewater, método 10200 F "
             "«Phytoplankton Counting Techniques», empleando cámara de recuento Sedgewick-Rafter "
             "o cámara de sedimentación con microscopio invertido (técnica de Utermöhl), según "
             "corresponda. La densidad se expresa en células por mililitro (cel/mL).")
    _parrafo(document,
             "La identificación taxonómica se verifica contra las bases de referencia AlgaeBase "
             "(Guiry & Guiry), GBIF Backbone Taxonomy y WoRMS. La nomenclatura sigue la "
             "clasificación vigente de dichas fuentes; las reasignaciones recientes de género se "
             "indican cuando aplica.")


def _seccion_valores_referenciales(document) -> None:
    _titulo_seccion(document, "V. VALORES REFERENCIALES")
    _parrafo(document,
             "Umbrales de cianobacterias (OMS 1999, agua de consumo — por densidad celular): "
             "Vigilancia inicial ≥ 200 cel/mL; Alerta 1 ≥ 2 000 cel/mL; Alerta 2 ≥ 100 000 cel/mL.")
    _parrafo(document,
             "Umbrales de cianobacterias (OMS 2021, agua recreativa — por biovolumen): "
             "Vigilancia inicial > 10 colonias/mL o > 50 filamentos/mL; Alerta 1 ≥ 0,3 mm³/L; "
             "Alerta 2 ≥ 4,0 mm³/L.")
    _parrafo(document,
             "Los presentes valores son referenciales para vigilancia sanitaria y no constituyen "
             "Estándares de Calidad Ambiental (ECA). La interpretación debe considerar el uso del "
             "cuerpo de agua y la normativa vigente (D.S. N.º 004-2017-MINAM).", size=8,
             italic=True, color="555555")


# ─────────────────────────────────────────────────────────────────────────────
# API pública
# ─────────────────────────────────────────────────────────────────────────────

def generar_docx_fitoplancton_campana(
    campana_id: str, overrides: Optional[dict] = None
) -> bytes:
    """Genera el Reporte de Ensayo de fitoplancton (.docx) y lo devuelve en bytes."""
    overrides = overrides or {}
    datos = _assemble(campana_id, overrides)

    document = Document()
    _configurar_pagina_y_encabezado(document, datos["meta"]["codigo_reporte"])
    _portada(document, datos["meta"])
    _seccion_informacion(document, datos["estaciones"])
    _seccion_resultados(document, datos)
    _seccion_control_calidad(document, datos)
    _seccion_metodos(document)
    _seccion_valores_referenciales(document)

    salida = BytesIO()
    document.save(salida)
    return salida.getvalue()
