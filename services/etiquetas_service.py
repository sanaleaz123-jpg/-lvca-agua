"""
services/etiquetas_service.py
Genera un archivo Word con etiquetas de frascos de muestreo para una campaña.

Parte de la plantilla `ETIQUETAS_EnBlanco_10puntos.docx` (5 ensayos
predefinidos con su preservante correspondiente) y produce un .docx con
una hoja por cada combinación (punto, profundidad), distribuyendo las N
etiquetas de ensayo en una grilla de 2 columnas (3 izquierda + 2 derecha
para 5 ensayos) que reutiliza la outer table 3×3 de la plantilla.

Modos de muestreo:
    "superficial" → 1 hoja por punto, profundidad fija "0.3 m",
                    código = código del punto.
    "columna"     → 1 hoja por cada profundidad seleccionada (S/M/F)
                    de cada punto. Profundidad en blanco; código del
                    punto con sufijo "(S)", "(M)" o "(F)".

Funciones públicas:
    get_ensayos_disponibles()      → lista de los 5 ensayos de la plantilla
    generar_etiquetas_campana(...) → bytes del .docx generado
"""

from __future__ import annotations

import os
from copy import deepcopy
from datetime import datetime, date
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from database.client import get_db


# ─────────────────────────────────────────────────────────────────────────────
# Constantes
# ─────────────────────────────────────────────────────────────────────────────

# 5 ensayos fijos de la plantilla, con el preservante asociado (no editable —
# coinciden con las cajas pre-impresas en `ETIQUETAS_EnBlanco_10puntos.docx`).
# El color de FONDO del membrete (la franja del encabezado, que es una imagen)
# codifica el preservante: amarillo=HNO3, verde=Lugol, azul=S/P. El texto del
# ENSAYO va siempre en negro. Los ensayos sintéticos clonan una etiqueta base
# con su preservante (heredan su membrete); `membrete` permite sobreescribir la
# imagen del membrete por una variante propia (p. ej. lila para Zoo/Perifiton).
ENSAYOS_PLANTILLA: list[dict] = [
    {"nombre": "Hierro y manganeso disuelto", "preservante": "HNO3"},   # amarillo
    {"nombre": "Fitoplancton",                "preservante": "LUGOL"},  # verde
    {"nombre": "Clorofila A",                 "preservante": "S/P"},    # azul
    {"nombre": "Color",                       "preservante": "S/P"},    # azul
    {"nombre": "Fisicoquímicos y nutrientes", "preservante": "S/P"},    # azul
    # Ensayos añadidos por la plataforma (no pre-impresos): se sintetizan
    # clonando una etiqueta base con el MISMO preservante.
    {"nombre": "DBO5",              "preservante": "S/P",   "sintetico": True},  # azul (S/P)
    {"nombre": "Microcistina - LR", "preservante": "S/P",   "sintetico": True},  # azul (S/P)
    # Zooplancton y Perifiton: Lugol, pero con membrete lila para distinguirlos.
    {"nombre": "Zooplancton", "preservante": "LUGOL", "sintetico": True, "membrete": "membrete_lila.png"},
    {"nombre": "Perifiton",   "preservante": "LUGOL", "sintetico": True, "membrete": "membrete_lila.png"},
]

# Etiqueta base a clonar para cada preservante al sintetizar ensayos nuevos
# (debe ser una etiqueta pre-impresa con ese preservante marcado).
_BASE_SINTETICA = "Color"  # fallback (S/P)
_BASE_SINTETICA_POR_PRESERVANTE = {
    "S/P":   "Color",        # casilla S/P marcada, membrete azul
    "LUGOL": "Fitoplancton", # casilla LUGOL marcada, membrete verde
}

# Carpeta con las variantes de membrete (imágenes) referenciadas por `membrete`.
_MEMBRETE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")

# Mapeo de cada ensayo (frasco/etiqueta) a los códigos de parámetro de la
# cadena de custodia que cubre. Permite que la cadena marque "x" solo en los
# parámetros realmente analizados por punto/profundidad (ver
# services/cadena_custodia_service.py). Códigos en mayúsculas (P###).
ENSAYO_A_PARAMS_CADENA: dict[str, list[str]] = {
    "Color":                       ["P010"],                       # Color verdadero
    "DBO5":                        ["P019"],
    "Microcistina - LR":           ["P091"],                       # Microcistina LR
    "Fitoplancton":                ["P120"],
    "Clorofila A":                 ["P124"],
    "Zooplancton":                 ["P126"],
    "Perifiton":                   ["P130"],
    "Hierro y manganeso disuelto": ["P074", "P077"],               # Fe + Mn disuelto
    "Fisicoquímicos y nutrientes": [
        "P025", "P028", "P031", "P032", "P033", "P034",
        "P036", "P038", "P041", "P042", "P050",
    ],
}


def params_cadena_de_ensayos(ensayos) -> set[str]:
    """
    Códigos de parámetro de la cadena (P###, en mayúsculas) cubiertos por la
    lista de `ensayos` seleccionados. Ensayos sin mapeo se ignoran.
    """
    codigos: set[str] = set()
    for e in ensayos or []:
        codigos.update(ENSAYO_A_PARAMS_CADENA.get(e, []))
    return {c.upper() for c in codigos}

# Modos válidos de muestreo (parámetro `tipo_muestreo`).
MODO_SUPERFICIAL = "superficial"
MODO_COLUMNA     = "columna"

# Orden canónico de profundidades para muestreo en columna.
PROFUNDIDADES_COLUMNA = ["S", "M", "F"]

# Tipos de punto a los que aplica el muestreo por columna de agua (S/M/F):
# cuerpos lénticos profundos. El resto se muestrea solo en superficie.
TIPOS_COLUMNA = {"embalse", "laguna"}

# Mapeo tipo de punto → código de matriz (mismo criterio que
# cadena_custodia_service.py para mantener consistencia con la cadena oficial).
_TIPO_A_MATRIZ: dict[str, str] = {
    "laguna":    "ADL",
    "rio":       "ADR",
    "canal":     "ADR",
    "manantial": "AMA",
    "embalse":   "ADL",
    "pozo":      "ASUB",
}

# Texto pre-impreso de "MUESTREADO POR" en la plantilla — a reemplazar.
_MUESTREADO_POR_DEFAULT = "A. Llacho, A. Vilcapaza"

# Posiciones (fila, col) dentro de la outer table donde se ubican las etiquetas
# (col 1 es separadora). Orden de lectura: izquierda→derecha, arriba→abajo.
# La plantilla trae 3 filas (6 posiciones); cuando se requieren más etiquetas se
# clonan filas adicionales hasta 5 filas (10 posiciones) (ver _construir_hoja).
_POSICIONES_ETIQUETAS = [
    (0, 0), (0, 2),
    (1, 0), (1, 2),
    (2, 0), (2, 2),
    (3, 0), (3, 2),
    (4, 0), (4, 2),
]

# Ruta a la plantilla (raíz del proyecto LVCA).
_TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..", "ETIQUETAS_EnBlanco_10puntos.docx"
)


# ─────────────────────────────────────────────────────────────────────────────
# API pública
# ─────────────────────────────────────────────────────────────────────────────

def get_ensayos_disponibles() -> list[str]:
    """Nombres de los 5 ensayos disponibles en la plantilla, en orden."""
    return [e["nombre"] for e in ENSAYOS_PLANTILLA]


def generar_etiquetas_campana(
    campana_id:      str,
    responsables:    list[str],
    filas_seleccion: list[dict] | None = None,
) -> bytes:
    """
    Genera el .docx con etiquetas. Cada "fila de selección" produce una hoja
    auto-descrita (punto + profundidad) y lleva solo los ensayos marcados.

    Args:
        campana_id: UUID de la campaña.
        responsables: nombres de los responsables de campo. Se concatenan
            con coma para el campo "MUESTREADO POR".
        filas_seleccion: lista de renglones a generar. Cada renglón es un
            dict:
                {"punto_id": str,
                 "prof": "S" | "M" | "F" | None,   # None = agua superficial
                 "ensayos": ["DBO5", "Color", ...]}
            • prof None → CÓDIGO sin sufijo, PROF=0.3 m (agua superficial).
            • prof S/M/F → CÓDIGO con sufijo "(S)|(M)|(F)", PROF en blanco
              (columna de agua; solo aplica a embalse/laguna, lo decide la UI).
            Solo se genera hoja para renglones con al menos un ensayo.

    Returns:
        Bytes del archivo .docx listo para descarga.

    Raises:
        FileNotFoundError: si la plantilla no existe.
        ValueError: si la campaña no tiene puntos o si la selección no
            produce ninguna hoja (ningún ensayo marcado en ningún renglón).
    """
    if not os.path.exists(_TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Plantilla de etiquetas no encontrada: {_TEMPLATE_PATH}"
        )

    campana, puntos = _cargar_datos(campana_id)
    if not puntos:
        raise ValueError("La campaña no tiene puntos de muestreo vinculados.")

    fecha_str = _formatear_fecha(
        campana.get("fecha_inicio"), campana.get("fecha_fin")
    )
    resp_str  = (
        ", ".join(r.strip() for r in responsables if r and r.strip())
        or _MUESTREADO_POR_DEFAULT
    )

    slots = _construir_slots(puntos, filas_seleccion or [])
    if not slots:
        raise ValueError(
            "No hay ensayos marcados para ningún punto. Marca al menos un "
            "ensayo en algún punto (y profundidad, en embalses/lagunas)."
        )

    doc = Document(_TEMPLATE_PATH)
    etiqueta_templates = _extraer_etiqueta_templates(doc)
    outer_template     = _extraer_outer_template(doc)
    _limpiar_body(doc)

    body   = doc.element.body
    sectPr = body.find(qn("w:sectPr"))

    for i, slot in enumerate(slots):
        valores = {
            "ESTACION":       slot["estacion"],
            "CODIGO":         slot["codigo_etiqueta"],
            "FECHA":          fecha_str,
            "HORA":           "",
            "MATRIZ":         slot["matriz"],
            "PROF":           slot["prof_valor"],
            "MUESTREADO_POR": resp_str,
        }

        # Ensayos de este slot, reordenados según la plantilla.
        ensayos_slot = _ensayos_ordenados(slot["ensayos"])

        if i > 0:
            page_break = _crear_parrafo_salto_pagina()
            _anexar(body, sectPr, page_break)

        hoja = _construir_hoja(
            outer_template, etiqueta_templates, ensayos_slot, valores
        )
        _anexar(body, sectPr, hoja)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _ensayos_ordenados(nombres) -> list[dict]:
    """Filtra ENSAYOS_PLANTILLA a `nombres`, preservando el orden de la plantilla."""
    pedidos = set(nombres or [])
    return [e for e in ENSAYOS_PLANTILLA if e["nombre"] in pedidos]


# ─────────────────────────────────────────────────────────────────────────────
# Carga de datos
# ─────────────────────────────────────────────────────────────────────────────

def _cargar_datos(campana_id: str) -> tuple[dict, list[dict]]:
    """Trae la campaña y sus puntos ordenados por código."""
    db = get_db()

    camp = (
        db.table("campanas")
        .select("codigo, nombre, fecha_inicio, fecha_fin, frecuencia")
        .eq("id", campana_id)
        .single()
        .execute()
        .data
        or {}
    )

    pts_res = (
        db.table("campana_puntos")
        .select("puntos_muestreo(id, codigo, nombre, tipo)")
        .eq("campana_id", campana_id)
        .execute()
    )
    puntos = [
        r["puntos_muestreo"] for r in (pts_res.data or [])
        if r.get("puntos_muestreo")
    ]
    puntos = sorted(puntos, key=lambda p: p.get("codigo") or "")
    return camp, puntos


def _formatear_fecha(fecha_inicio, fecha_fin) -> str:
    """
    Devuelve la fecha a imprimir en la etiqueta según la duración:

      • Si la campaña dura un solo día (fecha_inicio == fecha_fin) →
        "DD/MM/AAAA" — la fecha está totalmente determinada.
      • Si dura varios días (o fecha_fin desconocida) →
        "___/MM/AAAA" — el día se llena a mano en campo.

    Devuelve cadena vacía si fecha_inicio no es interpretable.
    """
    f_ini = _to_date(fecha_inicio)
    if f_ini is None:
        return ""

    f_fin = _to_date(fecha_fin)
    if f_fin is not None and f_fin == f_ini:
        return f"{f_ini.day:02d}/{f_ini.month:02d}/{f_ini.year}"
    return f"___/{f_ini.month:02d}/{f_ini.year}"


def _to_date(valor):
    """Convierte un str ISO / date / datetime a date. None si no se puede."""
    if valor is None or valor == "":
        return None
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, date):
        return valor
    try:
        return datetime.fromisoformat(str(valor)[:10]).date()
    except (ValueError, TypeError):
        return None


def _construir_slots(
    puntos:          list[dict],
    filas_seleccion: list[dict],
) -> list[dict]:
    """
    Convierte cada fila de selección en un "slot" (1 hoja) con su propia lista
    de ensayos. Solo se crean slots con ≥1 ensayo. El orden respeta el de
    `filas_seleccion`.

    Cada fila: {"punto_id", "prof": "S"|"M"|"F"|None, "ensayos": [...]}.
      • prof None  → código sin sufijo, PROF=0.3 m (agua superficial).
      • prof S/M/F → código con sufijo "(S)|(M)|(F)", PROF en blanco (columna).
    """
    pts_por_id = {p.get("id"): p for p in puntos}
    slots: list[dict] = []
    for fila in filas_seleccion:
        pt = pts_por_id.get(fila.get("punto_id"))
        if pt is None:
            continue
        ensayos = fila.get("ensayos") or []
        if not ensayos:
            continue

        codigo = pt.get("codigo", "") or ""
        matriz = _TIPO_A_MATRIZ.get((pt.get("tipo") or "").lower(), "AN")
        prof   = fila.get("prof")

        if prof in PROFUNDIDADES_COLUMNA:
            codigo_etiqueta = f"{codigo} ({prof})"
            prof_valor      = ""
        else:
            codigo_etiqueta = codigo
            prof_valor      = "0.3 m"

        slots.append({
            "estacion":        pt.get("nombre", "") or "",
            "codigo_etiqueta": codigo_etiqueta,
            "matriz":          matriz,
            "prof_valor":      prof_valor,
            "ensayos":         ensayos,
        })
    return slots


# ─────────────────────────────────────────────────────────────────────────────
# Manipulación de la plantilla
# ─────────────────────────────────────────────────────────────────────────────

def _extraer_etiqueta_templates(doc) -> dict[str, object]:
    """
    Captura UNA etiqueta (elemento <w:tbl> outer del recuadro) por cada ensayo
    pre-impreso en la plantilla, recorriendo sus tablas externas, y sintetiza
    las etiquetas de los ensayos marcados `sintetico` (no pre-impresos) clonando
    una etiqueta S/P base. Retorna {nombre_ensayo → elemento XML clonable}.
    """
    # Solo los ensayos con caja pre-impresa se buscan en el XML de la plantilla.
    nombres_preimpresos = {
        e["nombre"] for e in ENSAYOS_PLANTILLA if not e.get("sintetico")
    }
    capturados: dict[str, object] = {}

    for outer_table in doc.tables:
        for row in outer_table.rows:
            for cell in row.cells:
                if not cell.tables:
                    continue
                etiqueta_outer = cell.tables[0]   # 1x1 wrapper con borde
                text_completo = etiqueta_outer._element.xml
                for nombre in nombres_preimpresos:
                    if nombre in capturados:
                        continue
                    if nombre in text_completo:
                        capturados[nombre] = etiqueta_outer._element
                        break
        if len(capturados) >= len(nombres_preimpresos):
            break

    # Sintetizar las etiquetas no pre-impresas (DBO5, Microcistina, Zooplancton,
    # Perifiton, …) clonando una etiqueta base con el MISMO preservante (hereda
    # su casilla de preservante y su membrete). Texto del ENSAYO en negro.
    for e in ENSAYOS_PLANTILLA:
        if not e.get("sintetico") or e["nombre"] in capturados:
            continue
        base_nombre = _BASE_SINTETICA_POR_PRESERVANTE.get(
            e.get("preservante"), _BASE_SINTETICA
        )
        base = capturados.get(base_nombre)
        if base is None:
            base = capturados.get(_BASE_SINTETICA)
        if base is None:
            continue
        capturados[e["nombre"]] = _sintetizar_etiqueta(base, e["nombre"])

    # Sobreescribir la imagen del membrete para los ensayos que lo definan
    # (p. ej. Zooplancton/Perifiton → membrete lila). Requiere el `doc` para
    # registrar la imagen y obtener su relación (rId).
    membrete_cache: dict[str, str] = {}
    for e in ENSAYOS_PLANTILLA:
        img = e.get("membrete")
        el = capturados.get(e["nombre"])
        if img and el is not None:
            _set_membrete(doc, el, img, membrete_cache)

    return capturados


def _celda_ensayo(etiqueta_element):
    """Celda <w:tc> con el valor del campo ENSAYO (fila 5, col 2), o None."""
    tbl = _encontrar_tabla_campos(etiqueta_element)
    if tbl is None:
        return None
    rows = tbl.findall(qn("w:tr"))
    if len(rows) < 6:
        return None
    tcs = rows[5].findall(qn("w:tc"))
    return tcs[2] if len(tcs) >= 3 else None


def _sintetizar_etiqueta(base_element, nombre: str):
    """
    Clona una etiqueta base (mismo preservante) y reescribe su campo ENSAYO
    (fila 5, celda 2) con `nombre`. El resto —incluida la casilla de preservante
    marcada y el membrete— se conserva del clon. Texto en negro (el de la base).
    """
    nuevo = deepcopy(base_element)
    tc = _celda_ensayo(nuevo)
    if tc is not None:
        _set_paragrafo_directo(tc, nombre)
    return nuevo


def _set_membrete(doc, etiqueta_element, image_filename: str, cache: dict) -> None:
    """
    Reapunta la imagen del membrete (franja del encabezado) de la etiqueta a
    `image_filename` (en _MEMBRETE_DIR), registrándola una sola vez en el `doc`.
    Si la imagen no existe, deja el membrete original sin tocar.
    """
    blip = etiqueta_element.find(".//" + qn("a:blip"))
    if blip is None:
        return
    rid = cache.get(image_filename)
    if rid is None:
        path = os.path.join(_MEMBRETE_DIR, image_filename)
        if not os.path.exists(path):
            return
        rid, _ = doc.part.get_or_add_image(path)
        cache[image_filename] = rid
    blip.set(qn("r:embed"), rid)


def _extraer_outer_template(doc):
    """
    Devuelve la primera outer table 3×3 de la plantilla (estructura de hoja
    con 2 columnas de etiquetas + columna separadora central).
    """
    return doc.tables[0]._element


def _limpiar_body(doc) -> None:
    """Quita todas las tablas y párrafos del body, preservando <w:sectPr>."""
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sectPr:
            continue
        body.remove(child)


def _anexar(body, sectPr, elemento) -> None:
    """Inserta `elemento` antes de <w:sectPr> o al final si no existe."""
    if sectPr is not None:
        sectPr.addprevious(elemento)
    else:
        body.append(elemento)


def _crear_parrafo_salto_pagina():
    """Párrafo con page break para separar slots (hojas)."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Construcción de la hoja
# ─────────────────────────────────────────────────────────────────────────────

def _construir_hoja(outer_template, etiqueta_templates, ensayos_a_incluir, valores):
    """
    Crea una outer table 3×3 (copia del template) y coloca las N etiquetas
    en las posiciones de lectura (izq→der, arriba→abajo). Las posiciones
    sobrantes quedan vacías.
    """
    hoja = deepcopy(outer_template)
    rows = hoja.findall(qn("w:tr"))

    # La plantilla trae 3 filas (6 posiciones). Si hacen falta más etiquetas,
    # clonar la última fila —con sus celdas vacías— para habilitar la fila 4
    # (posiciones (3,0) y (3,2)), permitiendo hasta 8 etiquetas por hoja.
    filas_necesarias = max(
        (ri for i, (ri, ci) in enumerate(_POSICIONES_ETIQUETAS)
         if i < len(ensayos_a_incluir)),
        default=0,
    ) + 1
    while len(rows) < filas_necesarias and rows:
        nueva_fila = deepcopy(rows[-1])
        for celda in nueva_fila.findall(qn("w:tc")):
            if celda.find(qn("w:tbl")) is not None:
                _vaciar_celda(celda)
        hoja.append(nueva_fila)
        rows = hoja.findall(qn("w:tr"))

    for i, (ri, ci) in enumerate(_POSICIONES_ETIQUETAS):
        if ri >= len(rows):
            continue
        celdas_fila = rows[ri].findall(qn("w:tc"))
        if ci >= len(celdas_fila):
            continue
        celda = celdas_fila[ci]

        if i < len(ensayos_a_incluir):
            ensayo = ensayos_a_incluir[i]
            tbl_template = etiqueta_templates.get(ensayo["nombre"])
            if tbl_template is None:
                _vaciar_celda(celda)
                continue
            etiq = deepcopy(tbl_template)
            _rellenar_etiqueta(etiq, valores)
            _reemplazar_contenido_celda(celda, etiq)
        else:
            _vaciar_celda(celda)

    return hoja


def _reemplazar_contenido_celda(celda, etiqueta_tbl) -> None:
    """Quita tablas/párrafos de la celda (conserva <w:tcPr>) y mete la etiqueta."""
    tcPr = qn("w:tcPr")
    for child in list(celda):
        if child.tag != tcPr:
            celda.remove(child)
    celda.append(etiqueta_tbl)
    # Word requiere un párrafo después de una tabla anidada en celda.
    celda.append(OxmlElement("w:p"))


def _vaciar_celda(celda) -> None:
    """Deja la celda con solo <w:tcPr> + un párrafo vacío."""
    tcPr = qn("w:tcPr")
    for child in list(celda):
        if child.tag != tcPr:
            celda.remove(child)
    celda.append(OxmlElement("w:p"))


# ─────────────────────────────────────────────────────────────────────────────
# Llenado de campos dentro de una etiqueta
# ─────────────────────────────────────────────────────────────────────────────

def _rellenar_etiqueta(etiqueta_outer, valores: dict[str, str]) -> None:
    """
    Rellena los campos de una etiqueta clonada.

    Estructura real de la tabla de campos (8 filas × 4 columnas, anidada):
      Fila 0 → logo (gridSpan=4)
      Fila 1 → ESTACIÓN | : | (valor) | PRESERVANTE+opciones (vMerge)
      Fila 2 → CÓDIGO   | : | (valor) | (vMerge)
      Fila 3 → FECHA    | : | (tabla 1×4: FECHA-val | HORA | : | HORA-val)
      Fila 4 → MATRIZ   | : | (tabla 1×4: MATRIZ-val | PROF. | : | PROF-val)
      Fila 5 → ENSAYO   | : | (valor pre-impreso, NO tocar)
      Fila 6 → MUESTREADO POR: (gridSpan=4, solo etiqueta)
      Fila 7 → valor MUESTREADO POR (gridSpan=4)

    Campos esperados en `valores`:
        ESTACION, CODIGO, FECHA, HORA, MATRIZ, PROF, MUESTREADO_POR
    """
    content = _encontrar_tabla_campos(etiqueta_outer)
    if content is None:
        return

    rows = content.findall(qn("w:tr"))
    if len(rows) < 8:
        return

    _setear_paragrafo_de_celda(rows[1], col_idx=2, texto=valores.get("ESTACION", ""))
    _setear_paragrafo_de_celda(rows[2], col_idx=2, texto=valores.get("CODIGO", ""))
    _setear_par_anidado(
        rows[3], col_idx=2,
        izq=valores.get("FECHA", ""),
        der=valores.get("HORA", ""),
    )
    _setear_par_anidado(
        rows[4], col_idx=2,
        izq=valores.get("MATRIZ", ""),
        der=valores.get("PROF", ""),
    )
    # Fila 5 (ENSAYO): pre-impreso por etiqueta — no tocar.
    # Fila 6 (label MUESTREADO POR): no tocar.
    _setear_paragrafo_de_celda(rows[7], col_idx=0, texto=valores.get("MUESTREADO_POR", ""))


def _encontrar_tabla_campos(etiqueta_outer):
    """Devuelve el <w:tbl> de 8 filas con los campos (None si no existe)."""
    for tbl in etiqueta_outer.iter(qn("w:tbl")):
        if len(tbl.findall(qn("w:tr"))) == 8:
            return tbl
    return None


def _setear_paragrafo_de_celda(tr, col_idx: int, texto: str) -> None:
    """
    Pone `texto` en el primer <w:r>/<w:t> del primer <w:p> directo de la
    celda en la posición indicada. No toca tablas anidadas ni párrafos
    adicionales.
    """
    tcs = tr.findall(qn("w:tc"))
    if len(tcs) <= col_idx:
        return
    _set_paragrafo_directo(tcs[col_idx], texto)


def _setear_par_anidado(tr, col_idx: int, izq: str, der: str) -> None:
    """
    Para las filas FECHA y MATRIZ: la celda en `col_idx` contiene una
    tabla anidada 1×4 cuyas celdas son [valor_izq, label_der, ':', valor_der].
    Coloca `izq` en la celda 0 y `der` en la celda 3 (preserva labels).
    """
    tcs = tr.findall(qn("w:tc"))
    if len(tcs) <= col_idx:
        return
    contenedor = tcs[col_idx]
    tbl_anidada = contenedor.find(qn("w:tbl"))
    if tbl_anidada is None:
        return
    fila_anidada = tbl_anidada.find(qn("w:tr"))
    if fila_anidada is None:
        return
    sub_celdas = fila_anidada.findall(qn("w:tc"))
    if len(sub_celdas) >= 1:
        _set_paragrafo_directo(sub_celdas[0], izq)
    if len(sub_celdas) >= 4:
        _set_paragrafo_directo(sub_celdas[3], der)


def _set_paragrafo_directo(tc, texto: str) -> None:
    """
    Modifica el primer <w:p> directo del <w:tc> (ignora tablas anidadas).
    Pone `texto` en el primer <w:r>/<w:t>. Si no hay <w:t>, crea uno.
    """
    parrafos = [child for child in tc if child.tag == qn("w:p")]
    if not parrafos:
        return
    p = parrafos[0]
    ts = list(p.iter(qn("w:t")))
    if ts:
        ts[0].text = texto
        ts[0].set(qn("xml:space"), "preserve")
        for t in ts[1:]:
            t.text = ""
        return
    # Sin <w:t>: crear un run con texto.
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = texto
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
