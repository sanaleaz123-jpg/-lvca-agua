"""
services/ficha_campo_service.py
Genera las Fichas de Identificación del Punto de Monitoreo en DOCX.

Usa el template Word de referencia (static/templates/ficha_campo_template.docx)
y genera un documento con una ficha por cada muestra de la campaña.

Funciones públicas:
    get_datos_fichas_campana(campana_id)        → lista de dicts
    generar_docx_fichas(campana_id)             → bytes (.docx)
"""

from __future__ import annotations

import copy
from io import BytesIO
from pathlib import Path

from database.client import get_admin_client
from services.parametro_registry import get_cat_params, get_columnas_parametros
from services.storage_service import get_croquis_url, get_fotos_campo, download_imagen

TEMPLATE_PATH = Path(__file__).parent.parent / "static" / "templates" / "ficha_campo_template.docx"


# ─────────────────────────────────────────────────────────────────────────────
# Datos de ejemplo por punto (mientras no estén en la BD)
# ─────────────────────────────────────────────────────────────────────────────
DATOS_PUNTO_DEFAULT = {
    "132EABla3": {
        "departamento": "AREQUIPA", "provincia": "AREQUIPA",
        "distrito": "SAN JUAN DE TARUCANI",
        "accesibilidad": "Aproximadamente a 200 metros (5 minutos) del campamento de AUTODEMA en Aguada Blanca, acceso por trocha carrozable.",
        "representatividad": "Caracteriza la calidad de la masa de agua lacustre adyacente a la compuerta de descarga de la represa Aguada Blanca.",
    },
    "134EBamp3E": {
        "departamento": "AREQUIPA", "provincia": "CAYLLOMA",
        "distrito": "CALLALLI",
        "accesibilidad": "A 15 km del desvío de la carretera Arequipa-Chivay, acceso por trocha carrozable hasta la represa Bamputañe.",
        "representatividad": "Caracteriza la calidad del agua embalsada en la represa Bamputañe, zona de regulación del sistema Chili.",
    },
    "132EChal3E": {
        "departamento": "AREQUIPA", "provincia": "CAYLLOMA",
        "distrito": "YANQUE",
        "accesibilidad": "Acceso por carretera asfaltada Arequipa-Chivay hasta el km 135, luego desvío por trocha 2 km.",
        "representatividad": "Caracteriza la calidad del agua del embalse Chalhuanca, fuente de regulación hídrica.",
    },
    "134ECond3": {
        "departamento": "AREQUIPA", "provincia": "CAYLLOMA",
        "distrito": "CALLALLI",
        "accesibilidad": "A 180 km de Arequipa por carretera asfaltada hasta Condoroma, luego 3 km por trocha.",
        "representatividad": "Caracteriza la calidad del agua de la represa Condoroma, principal regulador del sistema Chili.",
    },
    "134EDesp3E": {
        "departamento": "AREQUIPA", "provincia": "CAYLLOMA",
        "distrito": "CALLALLI",
        "accesibilidad": "Acceso por trocha carrozable desde la carretera Condoroma-Sibayo, a 5 km del poblado.",
        "representatividad": "Caracteriza la calidad del agua del embalse Dique de los Españoles.",
    },
    "132EFray3E": {
        "departamento": "AREQUIPA", "provincia": "AREQUIPA",
        "distrito": "POLOBAYA",
        "accesibilidad": "A 80 km de Arequipa por carretera asfaltada hasta Polobaya, luego 15 km por trocha.",
        "representatividad": "Caracteriza la calidad del agua del embalse El Frayle, regulador del sistema Chili.",
    },
    "134Pañe3E": {
        "departamento": "AREQUIPA", "provincia": "CAYLLOMA",
        "distrito": "SANTA LUCIA",
        "accesibilidad": "Acceso por carretera asfaltada Arequipa-Juliaca hasta el km 160, desvío a la laguna Pañe.",
        "representatividad": "Caracteriza la calidad del agua de la laguna El Pañe, captación del sistema Colca.",
    },
    "132EPill3E": {
        "departamento": "AREQUIPA", "provincia": "AREQUIPA",
        "distrito": "SAN ANTONIO DE CHUCA",
        "accesibilidad": "A 120 km de Arequipa por carretera asfaltada hasta Pillones, acceso directo.",
        "representatividad": "Caracteriza la calidad del agua del embalse Pillones.",
    },
    "132RSumb4": {
        "departamento": "AREQUIPA", "provincia": "AREQUIPA",
        "distrito": "SAN JUAN DE TARUCANI",
        "accesibilidad": "A 3 km del campamento AUTODEMA Aguada Blanca, acceso por trocha carrozable.",
        "representatividad": "Caracteriza la calidad del agua del río Sumbay, afluente de la cuenca del Chili.",
    },
    "132BBTuti3E": {
        "departamento": "AREQUIPA", "provincia": "CAYLLOMA",
        "distrito": "TUTI",
        "accesibilidad": "Acceso por carretera asfaltada Arequipa-Chivay, a 500 m de la bocatoma.",
        "representatividad": "Punto de captación de agua para el sistema de riego, bocatoma de Tuti.",
    },
    "132BDHuambo": {
        "departamento": "AREQUIPA", "provincia": "CAYLLOMA",
        "distrito": "HUAMBO",
        "accesibilidad": "Acceso por carretera asfaltada hasta Huambo, luego 1 km al desarenador.",
        "representatividad": "Desarenador del sistema de conducción, punto de sedimentación de partículas.",
    },
    "132BBPitay": {
        "departamento": "AREQUIPA", "provincia": "AREQUIPA",
        "distrito": "UCHUMAYO",
        "accesibilidad": "A 25 km de Arequipa por carretera asfaltada, acceso directo a la bocatoma.",
        "representatividad": "Bocatoma de Pitay, punto de captación del canal de riego Zamácola.",
    },
}


def _get_datos_punto(punto: dict) -> dict:
    """Obtiene datos de ficha del punto desde BD, con fallback a defaults."""
    codigo = punto.get("codigo", "")
    defaults = DATOS_PUNTO_DEFAULT.get(codigo, {})
    return {
        "departamento":     punto.get("departamento")     or defaults.get("departamento", "AREQUIPA"),
        "provincia":        punto.get("provincia")        or defaults.get("provincia", ""),
        "distrito":         punto.get("distrito")         or defaults.get("distrito", ""),
        "accesibilidad":    punto.get("accesibilidad")    or defaults.get("accesibilidad", ""),
        "representatividad":punto.get("representatividad")or defaults.get("representatividad", ""),
        "finalidad":        punto.get("finalidad")        or defaults.get("finalidad", "Monitoreo de vigilancia de calidad de agua."),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Carga de datos por campaña
# ─────────────────────────────────────────────────────────────────────────────

def get_datos_fichas_campana(campana_id: str) -> list[dict]:
    """Reúne datos de todas las muestras de una campaña para generar fichas."""
    db = get_admin_client()

    m_res = (
        db.table("muestras")
        .select(
            "id, codigo, fecha_muestreo, hora_recoleccion, clima, "
            "observaciones_campo, tipo_muestra, "
            "puntos_muestreo("
            "  id, codigo, nombre, descripcion, tipo, cuenca, subcuenca, "
            "  utm_este, utm_norte, utm_zona, latitud, longitud, altitud_msnm, "
            "  entidad_responsable, departamento, provincia, distrito, "
            "  accesibilidad, representatividad, finalidad, "
            "  ecas(id, codigo, nombre, categoria, subcategoria)"
            "), "
            "campanas(id, codigo, nombre, responsable_campo)"
        )
        .eq("campana_id", campana_id)
        .order("fecha_muestreo")
        .execute()
    )

    # Obtener mediciones in situ para todas las muestras
    muestra_ids = [m["id"] for m in (m_res.data or []) if m.get("id")]
    insitu_map: dict[str, dict] = {}
    if muestra_ids:
        i_res = (
            db.table("mediciones_insitu")
            .select("muestra_id, parametro, valor")
            .in_("muestra_id", muestra_ids)
            .execute()
        )
        for r in (i_res.data or []):
            mid = r["muestra_id"]
            if mid not in insitu_map:
                insitu_map[mid] = {}
            insitu_map[mid][r["parametro"]] = r["valor"]

    fichas = []
    for muestra in (m_res.data or []):
        punto = muestra.get("puntos_muestreo") or {}
        campana = muestra.get("campanas") or {}
        eca = punto.get("ecas") or {}
        datos_extra = _get_datos_punto(punto)

        punto_id = punto.get("id", "")
        muestra_id = muestra.get("id", "")
        croquis_url = get_croquis_url(punto_id) if punto_id else None
        fotos = get_fotos_campo(muestra_id) if muestra_id else []

        fichas.append({
            "muestra": muestra,
            "punto": {**punto, **datos_extra},
            "campana": campana,
            "eca": eca,
            "croquis_url": croquis_url,
            "foto_campo_url": fotos[0]["url"] if fotos else None,
            "insitu": insitu_map.get(muestra_id, {}),
        })

    return fichas


# ─────────────────────────────────────────────────────────────────────────────
# Helpers para manipulación del template DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _get_unique_cells(row):
    """Retorna celdas únicas de una fila (colapsando celdas fusionadas)."""
    seen = set()
    cells = []
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id not in seen:
            seen.add(cell_id)
            cells.append(cell)
    return cells


def _clear_cell_content(cell):
    """Elimina todo el contenido de una celda, manteniendo propiedades del párrafo."""
    from docx.oxml.ns import qn
    for p in cell.paragraphs:
        for child in list(p._element):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag != "pPr":
                p._element.remove(child)


def _set_cell(cell, text, bold=False):
    """Reemplaza el contenido de la celda con texto, formato Arial 10pt."""
    from docx.shared import Pt
    _clear_cell_content(cell)
    p = cell.paragraphs[0]
    run = p.add_run(str(text or ""))
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = bold


def _set_cell_labeled(cell, label, value):
    """Celda con label en negrita + valor normal."""
    from docx.shared import Pt
    _clear_cell_content(cell)
    p = cell.paragraphs[0]
    r1 = p.add_run(str(label or ""))
    r1.font.name = "Arial"
    r1.font.size = Pt(10)
    r1.bold = True
    r2 = p.add_run(str(value or ""))
    r2.font.name = "Arial"
    r2.font.size = Pt(10)


def _set_cell_centered(cell, text):
    """Celda con texto centrado."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _clear_cell_content(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text or ""))
    run.font.name = "Arial"
    run.font.size = Pt(10)


def _add_cell_image(cell, img_bytes, width_cm=7):
    """Agrega una imagen centrada en la celda."""
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _clear_cell_content(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(BytesIO(img_bytes), width=Cm(width_cm))


# ─────────────────────────────────────────────────────────────────────────────
# Llenado de la tabla template con datos de una ficha
# ─────────────────────────────────────────────────────────────────────────────
#
# Estructura del template (28 filas × 21 columnas con celdas fusionadas):
#   R0:  Título (21 cols merged, bg #9CC3E5)
#   R1:  "UBICACIÓN:" (21 cols merged)
#   R2:  Departamento(s4) | valor(s4) | Provincia(s3) | valor(s5) | Distrito(s3) | valor(s2)
#   R3:  "NOMBRE DE CAMPAÑA:"(s4) | valor(s17)
#   R4:  "DESCRIPCIÓN DEL PUNTO DE MONITOREO" (21 cols, bg #9CC3E5)
#   R5:  Cuenca(s2) | val(s4) | Descripción(s3) | val(s3) | Clasificación(s5) | val(s4)
#   R6:  Código(s2) | val(s4) | Fecha/hora(s3) | val(s3) | Responsables(s5) | val(s4)
#   R7:  ACCESIBILIDAD: texto (21 cols merged)
#   R8:  REPRESENTATIVIDAD: texto (21 cols merged)
#   R9:  FINALIDAD: texto (21 cols merged)
#   R10: "COORDENADAS — UTM (WGS 84):" (21 cols, bg #9CC3E5)
#   R11: Zona(s1)|val(s2)|Este(s2)|val(s3)|Norte(s3)|val(s3)|Altitud(s4)|val(s3)
#   R12: "MONITOREO" (21 cols, bg #9CC3E5)
#   R13: Bocatoma(s7)|check(s1)|Embalse(s5)|check(s2)|Río(s5)|check(s1)
#   R14: "PARÁMETROS DE MONITOREO" (21 cols, bg #9CC3E5)
#   R15: "Parámetros de Campo:"(s8) | "Parámetros Físico-químicos..."(s13)
#   R16–R25: Filas de parámetros con checks
#   R26: "FOTOGRAFÍA Y CROQUIS..." (21 cols, bg #9CC3E5)
#   R27: Imagen izq (s10) | Imagen der (s11)
#

def _fill_parametros(table, insitu: dict | None = None) -> None:
    """
    Llena las filas R16-R25 con los parámetros activos desde la BD.
    Si se proporcionan datos insitu, muestra los valores medidos junto
    a los parámetros de campo.

    Estructura del template (10 filas disponibles):
        R16–R21: 6 filas × 3 columnas de parámetros (campo, fisicoquímico, otro)
        R22: fila separadora / header hidrobiológicos
        R23–R25: 3 filas × 3 columnas

    Cada fila tiene 6 celdas únicas: [nombre, check, nombre, check, nombre, check]
    Se distribuyen los parámetros dinámicamente por categoría.
    """
    from services.parametro_registry import get_parametros_insitu as _get_insitu

    def uc(idx):
        return _get_unique_cells(table.rows[idx])

    # Obtener parámetros agrupados por categoría desde la BD
    cat_params = get_cat_params()
    columnas = dict(get_columnas_parametros())  # {codigo: label}

    # Preparar listas de nombres por categoría
    campo_codigos = cat_params.get("Campo (In situ)", [])
    campo = [columnas.get(c, c) for c in campo_codigos]
    fisico = [columnas.get(c, c) for c in cat_params.get("Fisicoquímico", [])]
    hidro = [columnas.get(c, c) for c in cat_params.get("Hidrobiológico", [])]

    # Mapear códigos de campo a claves insitu para mostrar valores
    insitu_vals = insitu or {}
    insitu_params = {p["codigo"]: p["clave"] for p in _get_insitu()}

    # ── R16–R21: 6 filas con 3 columnas ─────────────────────────────────
    n_rows_main = 6

    # Dividir fisicoquímicos en dos columnas
    fisico_col2 = fisico[:n_rows_main]
    fisico_col3 = fisico[n_rows_main:n_rows_main * 2]

    for i in range(n_rows_main):
        cells = uc(16 + i)
        # Columna 1: Campo (con valor insitu si disponible)
        if i < len(campo):
            codigo_campo = campo_codigos[i] if i < len(campo_codigos) else ""
            clave_insitu = insitu_params.get(codigo_campo, "")
            val_insitu = insitu_vals.get(clave_insitu)
            label = campo[i]
            if val_insitu is not None:
                label = f"{campo[i]}: {val_insitu}"
            _set_cell(cells[0], label)
            _set_cell_centered(cells[1], "X")
        else:
            _set_cell(cells[0], "")
            _set_cell_centered(cells[1], "")
        # Columna 2: Fisicoquímico primera parte
        if i < len(fisico_col2):
            _set_cell(cells[2], fisico_col2[i])
            _set_cell_centered(cells[3], "X")
        else:
            _set_cell(cells[2], "")
            _set_cell_centered(cells[3], "")
        # Columna 3: Fisicoquímico segunda parte
        if i < len(fisico_col3):
            _set_cell(cells[4], fisico_col3[i])
            _set_cell_centered(cells[5], "X")
        else:
            _set_cell(cells[4], "")
            _set_cell_centered(cells[5], "")

    # Fisicoquímicos restantes (si hay más de 12)
    fisico_restantes = fisico[n_rows_main * 2:]

    # ── R22: Header hidrobiológicos + primer fisicoquímico restante ───────
    cells_22 = uc(22)
    _set_cell(cells_22[0], "Parámetros hidrobiológicos:", bold=True)
    if fisico_restantes:
        _set_cell(cells_22[1], fisico_restantes[0])
        _set_cell_centered(cells_22[2], "X")
    else:
        _set_cell(cells_22[1], "")
        _set_cell_centered(cells_22[2], "")
    if len(cells_22) > 3:
        _set_cell(cells_22[3], "")
    if len(cells_22) > 4:
        _set_cell(cells_22[4], "")

    # ── R23–R25: Hidrobiológicos + fisicoquímicos restantes ──────────────
    fisico_extra = fisico_restantes[1:]

    for i in range(3):
        cells = uc(23 + i)
        # Columna 1: Hidrobiológico
        if i < len(hidro):
            _set_cell(cells[0], hidro[i])
            _set_cell_centered(cells[1], "X")
        else:
            _set_cell(cells[0], "")
            _set_cell_centered(cells[1], "")
        # Columna 2: Fisicoquímicos restantes
        if i < len(fisico_extra):
            _set_cell(cells[2], fisico_extra[i])
            _set_cell_centered(cells[3], "X")
        else:
            _set_cell(cells[2], "")
            _set_cell_centered(cells[3], "")
        # Columna 3: vacío
        _set_cell(cells[4], "")
        _set_cell(cells[5], "")


def _fill_ficha_table(table, datos: dict) -> None:
    """Llena una tabla del template con los datos de una ficha."""
    punto = datos["punto"]
    muestra = datos["muestra"]
    campana = datos["campana"]
    eca = datos["eca"]

    def uc(row_idx):
        return _get_unique_cells(table.rows[row_idx])

    # R2: Departamento, Provincia, Distrito (celdas de valor en índices 1, 3, 5)
    _set_cell(uc(2)[1], punto.get("departamento", ""))
    _set_cell(uc(2)[3], punto.get("provincia", ""))
    _set_cell(uc(2)[5], punto.get("distrito", ""))

    # R3: Nombre de campaña (celda de valor en índice 1)
    _set_cell(uc(3)[1], campana.get("nombre", ""))

    # R5: Cuenca, Descripción, Clasificación
    _set_cell(uc(5)[1], punto.get("cuenca", ""))
    _set_cell(uc(5)[3], punto.get("descripcion") or punto.get("nombre", ""))
    cat = eca.get("categoria", "")
    sub = eca.get("subcategoria", "")
    clasif = f"Categoría {cat} Subcategoría {sub}" if cat else ""
    _set_cell(uc(5)[5], clasif)

    # R6: Código, Fecha/hora, Responsables
    _set_cell(uc(6)[1], punto.get("codigo", ""))
    fecha = str(muestra.get("fecha_muestreo", ""))[:10]
    hora = muestra.get("hora_recoleccion") or ""
    _set_cell(uc(6)[3], f"{fecha}\n{hora}".strip() if fecha else "")
    _set_cell(uc(6)[5], campana.get("responsable_campo", ""))

    # R7–R9: Accesibilidad, Representatividad, Finalidad
    _set_cell_labeled(uc(7)[0], "ACCESIBILIDAD: ", punto.get("accesibilidad", ""))
    _set_cell_labeled(uc(8)[0], "REPRESENTATIVIDAD: ", punto.get("representatividad", ""))
    _set_cell_labeled(uc(9)[0], "FINALIDAD: ", "Monitoreo de vigilancia de calidad de agua.")

    # R11: Zona, Este, Norte, Altitud
    _set_cell(uc(11)[1], punto.get("utm_zona") or "19K")
    utm_e = punto.get("utm_este")
    _set_cell(uc(11)[3], f"{utm_e:.0f}" if utm_e else "")
    utm_n = punto.get("utm_norte")
    _set_cell(uc(11)[5], f"{utm_n:.0f}" if utm_n else "")
    alt = punto.get("altitud_msnm")
    _set_cell(uc(11)[7], f"{alt:.0f} msnm" if alt else "")

    # R13: Tipo de monitoreo (checks en índices 1, 3, 5)
    tipo = (punto.get("tipo") or "").lower()
    _set_cell(uc(13)[1], "X" if tipo in ("bocatoma", "desarenador", "canal") else "")
    _set_cell(uc(13)[3], "X" if tipo in ("embalse", "laguna") else "")
    _set_cell(uc(13)[5], "X" if tipo in ("rio", "manantial", "otro") else "")

    # R16–R25: Parámetros de monitoreo — llenado programático con datos insitu
    _fill_parametros(table, insitu=datos.get("insitu"))

    # R27: Imágenes (croquis a la izquierda, foto de campo a la derecha)
    cells_27 = uc(27)

    if datos.get("croquis_url"):
        img_bytes = download_imagen(datos["croquis_url"])
        if img_bytes:
            _add_cell_image(cells_27[0], img_bytes, width_cm=7)
        else:
            _set_cell_centered(cells_27[0], "(Sin croquis cargado)")
    else:
        _set_cell_centered(cells_27[0], "(Sin croquis cargado)")

    if datos.get("foto_campo_url"):
        img_bytes = download_imagen(datos["foto_campo_url"])
        if img_bytes:
            _add_cell_image(cells_27[1], img_bytes, width_cm=7)
        else:
            _set_cell_centered(cells_27[1], "(Sin foto de campo)")
    else:
        _set_cell_centered(cells_27[1], "(Sin foto de campo)")


# ─────────────────────────────────────────────────────────────────────────────
# Generación DOCX — todas las fichas de una campaña en un solo documento
# ─────────────────────────────────────────────────────────────────────────────

def generar_docx_fichas(campana_id: str) -> bytes:
    """Genera DOCX con todas las fichas de una campaña usando el template."""
    from docx import Document
    from docx.enum.text import WD_BREAK

    fichas = get_datos_fichas_campana(campana_id)
    if not fichas:
        raise ValueError("No hay muestras en esta campaña.")

    doc = Document(str(TEMPLATE_PATH))

    # Limpiar párrafos del template (ej. "ANEXOS ...")
    for p in doc.paragraphs:
        for run in p.runs:
            run.text = ""

    # Limpiar imágenes de ejemplo del template en R27 antes de clonar
    template_table = doc.tables[0]
    cells_27 = _get_unique_cells(template_table.rows[27])
    _clear_cell_content(cells_27[0])
    _clear_cell_content(cells_27[1])

    # Guardar XML limpio de la tabla para clonar fichas adicionales
    clean_tbl = copy.deepcopy(template_table._tbl)

    # Llenar primera ficha
    _fill_ficha_table(template_table, fichas[0])

    # Fichas adicionales
    for ficha in fichas[1:]:
        # Salto de página
        p_break = doc.add_paragraph()
        run = p_break.add_run()
        run.add_break(WD_BREAK.PAGE)

        # Clonar tabla e insertarla después del salto de página
        new_tbl = copy.deepcopy(clean_tbl)
        p_break._element.addnext(new_tbl)

        # Llenar la tabla clonada
        _fill_ficha_table(doc.tables[-1], ficha)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
